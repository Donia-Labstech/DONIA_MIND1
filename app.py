"""
DONIA MIND 5 — المعلم الذكي (DONIA SMART TEACHER) — v5.0 MILITARY-GRADE
═══════════════════════════════════════════════════════════════════════════
FULLY REPAIRED: call_llm defined, Arcee handshake fixed, PDF fonts working,
Save-to-RAG implemented, download matrix restored.
"""
import streamlit as st
import os
import sqlite3
import re
import json
import io
import uuid
from datetime import datetime
from dotenv import load_dotenv
from langchain_groq import ChatGroq
import plotly.express as px
import plotly.graph_objects as go
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
import pandas as pd
from PIL import Image
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter

# ========== PDF & ARABIC IMPORTS ==========
from fpdf import FPDF
import arabic_reshaper
from bidi.algorithm import get_display
try:
    from streamlit_mic_recorder import mic_recorder
    MIC_AVAILABLE = True
except ImportError:
    MIC_AVAILABLE = False
try:
    from tavily import TavilyClient
    TAVILY_AVAILABLE = True
except ImportError:
    TAVILY_AVAILABLE = False
import tempfile
import requests

# DOCX support
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

# OCR support
try:
    import pytesseract
    _TESSERACT_AVAILABLE = True
except ImportError:
    _TESSERACT_AVAILABLE = False

# PDF text extraction
try:
    import pdfplumber
    _PDFPLUMBER_AVAILABLE = True
except ImportError:
    _PDFPLUMBER_AVAILABLE = False

# Arcee integration
try:
    from arcee import Arcee
    _ARCEE_AVAILABLE = True
except ImportError:
    _ARCEE_AVAILABLE = False

load_dotenv()

# ═══════════════════════════════════════════════════════════
# FIX 9047: Unified LLM caller - DEFINED AT TOP LEVEL
# ═══════════════════════════════════════════════════════════
def call_llm(llm, prompt: str) -> str:
    """Safely call LangChain LLM and return content."""
    try:
        response = llm.invoke(prompt)
        return response.content
    except Exception as e:
        st.error(f"LLM call failed: {e}")
        return ""

def get_llm(model_name: str, api_key: str):
    return ChatGroq(model_name=model_name, groq_api_key=api_key, temperature=0.7)

# ═══════════════════════════════════════════════════════════
# FIX R3/9048: API key helper (supports st.secrets + env)
# ═══════════════════════════════════════════════════════════
def _get_api_key(key_name: str) -> str:
    try:
        if hasattr(st, "secrets") and st.secrets:
            if key_name in st.secrets:
                return str(st.secrets[key_name]).strip()
    except Exception:
        pass
    return os.getenv(key_name, "").strip()

DEFAULT_GROQ_MODEL = os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile")
GROQ_API_KEY = _get_api_key("GROQ_API_KEY")
ARCEE_API_KEY = _get_api_key("ARCEE_API_KEY")
TAVILY_API_KEY = _get_api_key("TAVILY_API_KEY")

COPYRIGHT_FOOTER_AR = "جميع حقوق الملكية محفوظة حصرياً لمختبر DONIA LABS TECH © 2026"
WELCOME_MESSAGE_AR = "أهلاً بك أستاذنا القدير في رحاب DONIA MIND v5.0.. معاً نصنع مستقبل التعليم الجزائري بذكاء واحترافية."

# Social URLs
SOCIAL_URL_WHATSAPP = os.getenv("DONIA_URL_WHATSAPP", "https://wa.me/213674661737")
SOCIAL_URL_LINKEDIN = os.getenv("DONIA_URL_LINKEDIN", "https://www.linkedin.com/in/donia-labs-tech-smart-ideas-lab")
SOCIAL_URL_FACEBOOK = os.getenv("DONIA_URL_FACEBOOK", "https://www.facebook.com/share/1An6GhVd56/")
SOCIAL_URL_TELEGRAM = os.getenv("DONIA_URL_TELEGRAM", "https://t.me/+LxRzVAK12HZmNTQ8")
APP_URL = os.getenv("DONIA_APP_URL", "https://doniamind1-pvnmwp3kdthtlfct7uhopm.streamlit.app/")

# ═══════════════════════════════════════════════════════════
# LaTeX cleaning & Arabic processing
# ═══════════════════════════════════════════════════════════
def clean_latex(text: str) -> str:
    text = re.sub(r'\\(?=\s)', '', text)
    text = re.sub(r'\$\$([^\$]+?)\$\$', r'$$\1$$', text)
    brace_count = text.count('{') - text.count('}')
    if brace_count > 0:
        text += '}' * brace_count
    text = re.sub(r'\\begin\{align\*\?}', r'\\begin{align*}', text)
    text = re.sub(r'\\end\{align\*\?}', r'\\end{align*}', text)
    return text

def reshape_arabic(text: str) -> str:
    if not text:
        return ""
    try:
        reshaped = arabic_reshaper.reshape(text)
        bidi_text = get_display(reshaped)
        return bidi_text
    except Exception:
        return text

def get_pdf_mode_for_subject(subject: str):
    s = (subject or "").strip()
    if any(lang in s for lang in ["الإنجليزية", "Anglais"]):
        return False, "English"
    if any(lang in s for lang in ["الفرنسية", "Français"]):
        return False, "French"
    return True, "Arabic"

# ═══════════════════════════════════════════════════════════
# FIX R4: Robust FPDF2 with Arabic reshaping & font fallback
# ═══════════════════════════════════════════════════════════
class ArabicFPDF(FPDF):
    """
    FIXED v5.1: Robust Unicode font loader.
    Priority: Amiri (CDN) → system DejaVu → bundled DejaVu (CDN) → safe fallback.
    Never passes Arabic text to a non-Unicode font.
    """
    # System font search paths (Streamlit Cloud = Ubuntu 22)
    _SYSTEM_DEJAVU = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        "/usr/share/fonts/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/TTF/DejaVuSans.ttf",
    ]
    _AMIRI_URLS = [
        ("https://cdn.jsdelivr.net/gh/google/fonts/ofl/amiri/Amiri-Regular.ttf",
         "https://cdn.jsdelivr.net/gh/google/fonts/ofl/amiri/Amiri-Bold.ttf"),
        ("https://github.com/googlefonts/amiri/raw/main/fonts/ttf/Amiri-Regular.ttf",
         "https://github.com/googlefonts/amiri/raw/main/fonts/ttf/Amiri-Bold.ttf"),
    ]
    _DEJAVU_URLS = [
        "https://cdn.jsdelivr.net/gh/dejavu-fonts/dejavu-fonts/ttf/DejaVuSans.ttf",
        "https://github.com/dejavu-fonts/dejavu-fonts/raw/master/ttf/DejaVuSans.ttf",
    ]

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.use_amiri = False
        self._active_font = "Helvetica"
        self.has_unicode = False

        base_dir = os.path.dirname(os.path.abspath(__file__))
        font_dir = os.path.join(base_dir, "fonts")
        os.makedirs(font_dir, exist_ok=True)
        reg_path  = os.path.join(font_dir, "Amiri-Regular.ttf")
        bold_path = os.path.join(font_dir, "Amiri-Bold.ttf")
        deja_path = os.path.join(font_dir, "DejaVuSans.ttf")

        # ── Step 1: try to download Amiri ──
        if not os.path.exists(reg_path) or os.path.getsize(reg_path) < 100_000:
            for url_r, url_b in self._AMIRI_URLS:
                try:
                    r = requests.get(url_r, timeout=12)
                    if r.status_code == 200 and len(r.content) > 100_000:
                        open(reg_path, "wb").write(r.content)
                        rb = requests.get(url_b, timeout=12)
                        if rb.status_code == 200:
                            open(bold_path, "wb").write(rb.content)
                        break
                except Exception:
                    continue

        # ── Step 2: try to load Amiri ──
        if os.path.exists(reg_path) and os.path.getsize(reg_path) > 100_000:
            try:
                self.add_font("Amiri", "", reg_path)
                b_src = bold_path if os.path.exists(bold_path) and os.path.getsize(bold_path) > 100_000 else reg_path
                self.add_font("Amiri", "B", b_src)
                self.set_font("Amiri", size=12)
                self.use_amiri = True
                self._active_font = "Amiri"
                self.has_unicode = True
                return
            except Exception:
                pass

        # ── Step 3: try system DejaVu (often present on Ubuntu/Streamlit Cloud) ──
        for sys_path in self._SYSTEM_DEJAVU:
            if os.path.exists(sys_path):
                try:
                    self.add_font("DejaVu", "", sys_path)
                    self.set_font("DejaVu", size=12)
                    self._active_font = "DejaVu"
                    self.has_unicode = True
                    return
                except Exception:
                    continue

        # ── Step 4: download DejaVu ──
        if not os.path.exists(deja_path) or os.path.getsize(deja_path) < 200_000:
            for url in self._DEJAVU_URLS:
                try:
                    r = requests.get(url, timeout=12)
                    if r.status_code == 200 and len(r.content) > 200_000:
                        open(deja_path, "wb").write(r.content)
                        break
                except Exception:
                    continue
        if os.path.exists(deja_path) and os.path.getsize(deja_path) > 200_000:
            try:
                self.add_font("DejaVu", "", deja_path)
                self.set_font("DejaVu", size=12)
                self._active_font = "DejaVu"
                self.has_unicode = True
                return
            except Exception:
                pass

        # ── Step 5: last resort – Helvetica (ASCII only, Arabic will be skipped) ──
        self.set_font("Helvetica", size=12)
        self._active_font = "Helvetica"
        self.has_unicode = False

    def _safe_text(self, text: str, rtl: bool = True) -> str:
        """Return text safe for current font. If no Unicode font, strip non-ASCII."""
        if self.has_unicode:
            return reshape_arabic(text) if rtl else text
        # Fallback: keep only printable ASCII to avoid crash
        return re.sub(r'[^\x20-\x7E]', '?', text)

    def _font(self, bold: bool = False) -> str:
        if self._active_font == "Amiri":
            return "Amiri"
        if self._active_font == "DejaVu":
            return "DejaVu"
        return "Helvetica"

    def set_sized_font(self, size: int, bold: bool = False):
        style = "B" if bold and self._active_font == "Amiri" else ""
        self.set_font(self._font(), style, size)

    def multi_cell_text(self, text, w, align='R', rtl=True):
        safe = self._safe_text(text, rtl)
        self.multi_cell(w, 6, safe, border=0, align=align)

def ensure_font_files():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    font_dir = os.path.join(base_dir, "fonts")
    os.makedirs(font_dir, exist_ok=True)
    pairs = (
        ("Amiri-Regular.ttf", "https://github.com/googlefonts/amiri/raw/main/fonts/ttf/Amiri-Regular.ttf"),
        ("Amiri-Bold.ttf", "https://github.com/googlefonts/amiri/raw/main/fonts/ttf/Amiri-Bold.ttf"),
    )
    for fname, url in pairs:
        path = os.path.join(font_dir, fname)
        if not os.path.exists(path) or os.path.getsize(path) < 100000:
            try:
                r = requests.get(url, timeout=10)
                with open(path, "wb") as f:
                    f.write(r.content)
            except Exception:
                pass

# ========== PDF GENERATORS ==========
def generate_simple_pdf(content: str, title: str, subtitle: str = "", rtl: bool = True) -> bytes:
    ensure_font_files()
    pdf = ArabicFPDF()
    pdf.add_page()
    pdf.set_sized_font(14)
    if rtl:
        pdf.cell(0, 8, pdf._safe_text("الجمهورية الجزائرية الديمقراطية الشعبية"), ln=True, align='C')
        pdf.cell(0, 8, pdf._safe_text("وزارة التربية الوطنية"), ln=True, align='C')
        pdf.cell(0, 8, pdf._safe_text(f"DONIA MIND — {title}"), ln=True, align='C')
    else:
        pdf.cell(0, 8, "Algerian Democratic Republic", ln=True, align='C')
        pdf.cell(0, 8, "Ministry of Education", ln=True, align='C')
        pdf.cell(0, 8, f"DONIA MIND — {title}", ln=True, align='C')
    pdf.ln(5)
    pdf.set_sized_font(11)
    for line in content.splitlines():
        line = line.strip()
        if not line:
            pdf.ln(3)
            continue
        if line.startswith("##"):
            pdf.set_sized_font(12, bold=True)
            pdf.multi_cell_text(line[2:], 190, align='R' if rtl else 'L', rtl=rtl)
            pdf.set_sized_font(11)
        else:
            pdf.multi_cell_text(line, 190, align='R' if rtl else 'L', rtl=rtl)
        pdf.ln(2)
    pdf.set_y(-15)
    pdf.set_sized_font(8)
    pdf.cell(0, 10, pdf._safe_text(COPYRIGHT_FOOTER_AR) if rtl else COPYRIGHT_FOOTER_AR, align='C')
    return bytes(pdf.output())

def generate_exam_pdf(exam_data: dict) -> bytes:
    ensure_font_files()
    pdf = ArabicFPDF()
    pdf.add_page()
    subj = exam_data.get("subject", "")
    rtl, _ = get_pdf_mode_for_subject(subj)
    pdf.set_sized_font(10)
    pdf.cell(95, 8, pdf._safe_text(exam_data.get("school", ""), rtl), border=1, align='C')
    pdf.cell(95, 8,
             pdf._safe_text("الجمهورية الجزائرية الديمقراطية الشعبية", rtl) if rtl else "Algerian Republic",
             border=1, ln=True, align='C')
    pdf.cell(95, 8,
             pdf._safe_text(f"المستوى: {exam_data.get('grade', '')}", rtl) if rtl else f"Level: {exam_data.get('grade', '')}",
             border=1)
    pdf.cell(95, 8,
             pdf._safe_text(f"المدة: {exam_data.get('duration', '')}", rtl) if rtl else f"Duration: {exam_data.get('duration', '')}",
             border=1, ln=True)
    pdf.ln(8)
    pdf.set_sized_font(14, bold=True)
    title = (f"اختبار {exam_data.get('semester', '')} في مادة {subj}"
             if rtl else f"Exam — {exam_data.get('semester', '')} — {subj}")
    pdf.cell(0, 10, pdf._safe_text(title, rtl) if rtl else title, ln=True, align='C')
    pdf.ln(5)
    pdf.set_sized_font(11)
    content_text = exam_data.get("content", "")
    for line in content_text.splitlines():
        line = line.strip()
        if not line:
            pdf.ln(2)
            continue
        pdf.multi_cell_text(line, 190, align='R' if rtl else 'L', rtl=rtl)
        pdf.ln(1)
    pdf.set_y(-15)
    pdf.set_sized_font(8)
    pdf.cell(0, 10, pdf._safe_text(COPYRIGHT_FOOTER_AR, rtl), align='C')
    return bytes(pdf.output())
def generate_report_pdf(report_data: dict) -> bytes:
    ensure_font_files()
    pdf = ArabicFPDF()
    pdf.add_page()
    pdf.set_sized_font(12)
    pdf.cell(0, 10, pdf._safe_text("تحليل نتائج الأقسام"), ln=True, align='C')
    pdf.ln(5)
    for cls in report_data.get('classes', []):
        pdf.set_sized_font(12, bold=True)
        pdf.cell(0, 8, pdf._safe_text(f"القسم: {cls['name']}"), ln=True, align='R')
        pdf.set_sized_font(11)
        stats = f"عدد التلاميذ: {cls.get('total',0)}  |  المعدل: {cls.get('avg',0):.2f}  |  النجاح: {cls.get('pass_rate',0):.1f}%"
        pdf.multi_cell_text(stats, 190, align='R', rtl=True)
        pdf.ln(4)
        if cls.get('top5'):
            pdf.set_sized_font(11, bold=True)
            pdf.cell(0, 6, pdf._safe_text("أفضل 5 تلاميذ"), ln=True, align='R')
            pdf.set_sized_font(10)
            for i, s in enumerate(cls['top5'][:5], 1):
                pdf.cell(0, 5, pdf._safe_text(f"{i}. {s['name']} — {s['avg']:.2f}"), ln=True, align='R')
        pdf.ln(6)
    if report_data.get('ai_analysis'):
        pdf.set_sized_font(11, bold=True)
        pdf.cell(0, 8, pdf._safe_text("التحليل البيداغوجي"), ln=True, align='R')
        pdf.set_sized_font(10)
        for line in report_data['ai_analysis'].splitlines():
            if line.strip():
                pdf.multi_cell_text(line, 190, align='R', rtl=True)
    pdf.set_y(-15)
    pdf.set_sized_font(8)
    pdf.cell(0, 10, pdf._safe_text(COPYRIGHT_FOOTER_AR), align='C')
    return bytes(pdf.output())
def generate_lesson_plan_pdf(plan_data: dict) -> bytes:
    ensure_font_files()
    pdf = ArabicFPDF()
    pdf.add_page()
    pdf.set_sized_font(11)
    pdf.cell(0, 8, pdf._safe_text("الجمهورية الجزائرية الديمقراطية الشعبية"), ln=True, align='C')
    pdf.cell(0, 8, pdf._safe_text("وزارة التربية الوطنية"), ln=True, align='C')
    pdf.ln(4)
    pdf.set_sized_font(13, bold=True)
    header_text = (f"مذكرة رقم: ____ | المؤسسة: {plan_data.get('school','')} | "
                   f"الأستاذ(ة): {plan_data.get('teacher','')}")
    pdf.cell(0, 8, pdf._safe_text(header_text), ln=True, align='R')
    pdf.ln(5)
    pdf.set_sized_font(10)
    info = [
        ("الميدان",         plan_data.get('domain', '')),
        ("المستوى",         plan_data.get('grade', '')),
        ("الباب",           plan_data.get('chapter', '')),
        ("المدة",           plan_data.get('duration', '')),
        ("المورد المعرفي",  plan_data.get('lesson', '')),
        ("نوع الحصة",       plan_data.get('session_type', '')),
    ]
    for i in range(0, len(info), 2):
        pdf.cell(45, 7, pdf._safe_text(info[i][0]),   border=1)
        pdf.cell(50, 7, pdf._safe_text(info[i][1]),   border=1)
        pdf.cell(45, 7, pdf._safe_text(info[i+1][0]), border=1)
        pdf.cell(50, 7, pdf._safe_text(info[i+1][1]), border=1, ln=True)
    pdf.ln(5)
    for s in ["المرحلة", "المدة", "سير الدرس", "التقويم"]:
        pdf.cell(47.5, 7, pdf._safe_text(s), border=1)
    pdf.ln()
    rows = [
        ("تهيئة",            plan_data.get('duration_t', '5 د'),  plan_data.get('intro', ''),    plan_data.get('eval', '')),
        ("بناء الموارد",     plan_data.get('duration_b', '25 د'), plan_data.get('build', ''),    ""),
        ("إعادة الاستثمار", plan_data.get('duration_r', '15 د'), plan_data.get('reinvest', ''), ""),
        ("الواجب المنزلي",  "",                                    plan_data.get('homework', ''), ""),
    ]
    for row in rows:
        pdf.cell(47.5, 40, pdf._safe_text(row[0]), border=1)
        pdf.cell(47.5, 40, pdf._safe_text(row[1]), border=1)
        pdf.multi_cell(47.5, 5, pdf._safe_text(row[2]), border=1)
        pdf.cell(47.5, 40, pdf._safe_text(row[3]), border=1)
        pdf.ln()
    pdf.set_y(-15)
    pdf.set_sized_font(8)
    pdf.cell(0, 10, pdf._safe_text(COPYRIGHT_FOOTER_AR), align='C')
    return bytes(pdf.output())


# ========== DOCX generators ==========
def _docx_set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi_el = OxmlElement('w:bidi')
    bidi_el.set(qn('w:val'), '1')
    pPr.append(bidi_el)

def _docx_heading(doc, text: str, level: int = 1, color_hex: str = "145a32"):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _docx_set_rtl(p)
    for run in p.runs:
        r, g, b = (int(color_hex[i:i+2], 16) for i in (0, 2, 4))
        run.font.color.rgb = RGBColor(r, g, b)
    return p

def _docx_para(doc, text: str, bold: bool = False, size: int = 12, align=WD_ALIGN_PARAGRAPH.RIGHT):
    p = doc.add_paragraph()
    p.alignment = align
    _docx_set_rtl(p)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def generate_exam_docx(exam_data: dict) -> bytes:
    if not _DOCX_AVAILABLE:
        return b""
    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    hdr = doc.add_table(rows=3, cols=2)
    hdr.style = 'Table Grid'
    cells = hdr.rows[0].cells
    cells[0].text = exam_data.get('school', '')
    cells[1].text = "الجمهورية الجزائرية الديمقراطية الشعبية"
    cells = hdr.rows[1].cells
    cells[0].text = f"السنة الدراسية: {exam_data.get('year', '')}"
    cells[1].text = "وزارة التربية الوطنية"
    cells = hdr.rows[2].cells
    cells[0].text = f"المدة: {exam_data.get('duration', '')}"
    cells[1].text = (f"المستوى: {exam_data.get('grade', '')}   |   "
                     f"المقاطعة: {exam_data.get('district', '')}")
    for row in hdr.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _docx_set_rtl(para)
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _docx_set_rtl(title_p)
    run = title_p.add_run(f"اختبار {exam_data.get('semester', '')} في مادة {exam_data.get('subject', '')}")
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()
    content = exam_data.get('content', '')
    for line in content.split('\n'):
        _docx_para(doc, line)
    doc.add_paragraph()
    sign_p = doc.add_paragraph("بالتوفيق                                              إنتهى")
    sign_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

def generate_lesson_plan_docx(plan_data: dict) -> bytes:
    if not _DOCX_AVAILABLE:
        return b""
    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    _docx_heading(doc, "المذكرة البيداغوجية — DONIA MIND", level=1)
    info_rows = [
        ["المؤسسة", plan_data.get('school', '')],
        ["الأستاذ(ة)", plan_data.get('teacher', '')],
        ["المادة", plan_data.get('subject', '')],
        ["المستوى", plan_data.get('grade', '')],
        ["الدرس", plan_data.get('lesson', '')],
        ["المجال", plan_data.get('domain', '')],
        ["المدة الإجمالية", plan_data.get('duration', '')],
    ]
    tbl = doc.add_table(rows=len(info_rows), cols=2)
    tbl.style = 'Table Grid'
    for i, (label, val) in enumerate(info_rows):
        cells = tbl.rows[i].cells
        cells[0].text = label
        cells[1].text = val
        for cell in cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _docx_set_rtl(para)
    doc.add_paragraph()
    _docx_heading(doc, "محتوى المذكرة", level=2, color_hex="1e8449")
    content = plan_data.get('content', '')
    for line in content.split('\n'):
        _docx_para(doc, line)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

def generate_report_docx(report_data: dict) -> bytes:
    if not _DOCX_AVAILABLE:
        return b""
    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    _docx_heading(doc, "تقرير تحليل نتائج الأقسام", level=1)
    _docx_para(doc, f"المادة: {report_data.get('subject', '')}   |   "
               f"الفصل: {report_data.get('semester', '')}   |   "
               f"المؤسسة: {report_data.get('school', '')}", bold=True)
    doc.add_paragraph()
    for cls in report_data.get('classes', []):
        _docx_heading(doc, f"القسم: {cls.get('name', '')}", level=2, color_hex="1e8449")
        stats_rows = [
            ["عدد التلاميذ", str(cls.get('total', ''))],
            ["المعدل العام", str(cls.get('avg', ''))],
            ["أعلى معدل", str(cls.get('max', ''))],
            ["أدنى معدل", str(cls.get('min', ''))],
            ["نسبة النجاح %", str(cls.get('pass_rate', ''))],
        ]
        tbl = doc.add_table(rows=len(stats_rows), cols=2)
        tbl.style = 'Table Grid'
        for i, (label, val) in enumerate(stats_rows):
            cells = tbl.rows[i].cells
            cells[0].text = label
            cells[1].text = val
            for cell in cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    _docx_set_rtl(para)
        doc.add_paragraph()
        top5 = cls.get('top5', [])
        if top5:
            _docx_para(doc, "أفضل 5 تلاميذ:", bold=True)
            for idx, s in enumerate(top5, 1):
                _docx_para(doc, f"  {idx}. {s['name']} — {s['avg']:.2f}")
        doc.add_paragraph()
    if report_data.get('ai_analysis'):
        _docx_heading(doc, "التقرير البيداغوجي (الذكاء الاصطناعي)", level=2, color_hex="922b21")
        for line in report_data['ai_analysis'].split('\n'):
            _docx_para(doc, line)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# ═══════════════════════════════════════════════════════════
# FIX R3: Arcee real handshake & critic layer
# ═══════════════════════════════════════════════════════════
def get_arcee_client():
    if not _ARCEE_AVAILABLE or not ARCEE_API_KEY:
        return None
    try:
        return Arcee(api_key=ARCEE_API_KEY)
    except Exception as e:
        st.warning(f"Arcee init error: {e}")
        return None

def test_arcee_connection() -> bool:
    """
    FIXED v5.1: Try SDK first; fall back to direct HTTP REST probe.
    Returns True only if the API key is valid and the server is reachable.
    """
    if not ARCEE_API_KEY:
        return False
    # Path 1: official SDK
    if _ARCEE_AVAILABLE:
        try:
            client = get_arcee_client()
            if client is not None:
                if hasattr(client, "generate"):
                    client.generate("ping")
                    return True
                elif hasattr(client, "list_retrievers"):
                    client.list_retrievers()
                    return True
        except Exception:
            pass
    # Path 2: direct HTTP probe (works without SDK)
    for endpoint in [
        "https://models.arcee.ai/v1/models",
        "https://api.arcee.ai/v2/models",
        "https://api.arcee.ai/v1/models",
    ]:
        try:
            resp = requests.get(
                endpoint,
                headers={"Authorization": f"Bearer {ARCEE_API_KEY}"},
                timeout=7,
            )
            if resp.status_code == 200:
                return True
            if resp.status_code in (401, 403):
                return False   # Key rejected — server alive but key wrong
        except Exception:
            continue
    return False
def call_arcee_generate(prompt: str) -> str:
    if not _ARCEE_AVAILABLE or not ARCEE_API_KEY:
        raise Exception("Arcee not available")
    client = get_arcee_client()
    if client is None:
        raise Exception("Arcee client init failed")
    if hasattr(client, "generate"):
        return client.generate(prompt)
    else:
        raise Exception("Arcee client has no generate method")

def arcee_critic(content: str, subject: str, grade: str) -> dict:
    if not ARCEE_API_KEY or not _ARCEE_AVAILABLE:
        try:
            llm = get_llm(DEFAULT_GROQ_MODEL, GROQ_API_KEY)
            critic_prompt = f"""أنت ناقد بيداغوجي جزائري. حلل المحتوى التالي:
المادة: {subject}  |  المستوى: {grade}

المحتوى:
{content[:2000]}

أجب بصيغة JSON:
{{
    "aligned": true/false,
    "score": 0-10,
    "remarks": "نقاط الضعف أو التحسينات",
    "suggestions": "اقتراحات"
}}"""
            response = call_llm(llm, critic_prompt)
            try:
                return json.loads(response)
            except:
                return {"aligned": True, "score": 7, "remarks": "تعذر التحقق", "suggestions": ""}
        except:
            return {"aligned": True, "score": 7, "remarks": "معطل", "suggestions": ""}
    else:
        try:
            prompt = f"تحقق من مطابقة المحتوى لمنهاج {subject} المستوى {grade}. أجب بصيغة JSON: {{'aligned': bool, 'score': int, 'remarks': str, 'suggestions': str}}"
            result_text = call_arcee_generate(prompt)
            try:
                return json.loads(result_text)
            except:
                import re
                aligned_match = re.search(r'"aligned":\s*(true|false)', result_text, re.IGNORECASE)
                score_match = re.search(r'"score":\s*(\d+)', result_text)
                remarks_match = re.search(r'"remarks":\s*"([^"]*)"', result_text)
                suggestions_match = re.search(r'"suggestions":\s*"([^"]*)"', result_text)
                aligned = aligned_match and aligned_match.group(1).lower() == 'true'
                score = int(score_match.group(1)) if score_match else 7
                remarks = remarks_match.group(1) if remarks_match else "تم التحقق"
                suggestions = suggestions_match.group(1) if suggestions_match else ""
                return {"aligned": aligned, "score": score, "remarks": remarks, "suggestions": suggestions}
        except Exception as e:
            st.error(f"Arcee validation error: {e}")
            return {"aligned": True, "score": 7, "remarks": "خطأ في Arcee", "suggestions": ""}

def dual_llm_generate_with_critic(prompt: str, subject: str, grade: str, use_critic: bool = True) -> tuple[str, dict]:
    if not GROQ_API_KEY:
        return "", {"error": "GROQ_API_KEY missing"}
    llm = get_llm(DEFAULT_GROQ_MODEL, GROQ_API_KEY)
    generated = call_llm(llm, prompt)
    critic_report = {"validated": False, "original": generated[:300]}
    if use_critic:
        critic = arcee_critic(generated, subject, grade)
        critic_report["validated"] = critic.get("aligned", False)
        critic_report["score"] = critic.get("score", 0)
        critic_report["remarks"] = critic.get("remarks", "")
        critic_report["suggestions"] = critic.get("suggestions", "")
    return generated, critic_report

# ═══════════════════════════════════════════════════════════
# NEW: Save to RAG (store generated content in knowledge base)
# ═══════════════════════════════════════════════════════════
def save_to_rag(content: str, content_type: str, metadata: dict):
    """Store generated text into RAG database for future retrieval."""
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS rag_documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            content TEXT,
            content_type TEXT,
            metadata_json TEXT,
            created_at TEXT
        )
    """)
    conn.execute(
        "INSERT INTO rag_documents (content, content_type, metadata_json, created_at) VALUES (?,?,?,?)",
        (content, content_type, json.dumps(metadata, ensure_ascii=False), datetime.now().strftime("%Y-%m-%d %H:%M"))
    )
    conn.commit()
    conn.close()

# ═══════════════════════════════════════════════════════════
# FIX: Excel text export — generates a styled .xlsx from
#      any generated text content (lesson plan / exam / exercise)
# ═══════════════════════════════════════════════════════════
def generate_text_excel(content: str, title: str, metadata: dict) -> bytes:
    """Wrap generated text into a styled Excel workbook."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "المحتوى"
    ws.sheet_view.rightToLeft = True

    header_font  = Font(name="Arial", bold=True, size=13, color="FFFFFF")
    meta_font    = Font(name="Arial", bold=True, size=11, color="145A32")
    body_font    = Font(name="Arial", size=11)
    center       = Alignment(horizontal="center", vertical="center", wrap_text=True)
    right        = Alignment(horizontal="right",  vertical="top",    wrap_text=True)
    green_fill   = PatternFill("solid", fgColor="145A32")
    light_fill   = PatternFill("solid", fgColor="EAF6EE")
    border_side  = Side(style="thin", color="27AE60")
    cell_border  = Border(left=border_side, right=border_side,
                          top=border_side, bottom=border_side)

    # Row 1 – Main title
    ws.merge_cells("A1:C1")
    ws["A1"] = title
    ws["A1"].font = header_font
    ws["A1"].fill = green_fill
    ws["A1"].alignment = center
    ws.row_dimensions[1].height = 32

    # Rows 2-N – metadata key/value pairs
    row = 2
    for k, v in metadata.items():
        ws.cell(row=row, column=1, value=k).font = meta_font
        ws.cell(row=row, column=1).fill = light_fill
        ws.cell(row=row, column=1).alignment = right
        ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=3)
        ws.cell(row=row, column=2, value=str(v)).font = body_font
        ws.cell(row=row, column=2).alignment = right
        row += 1

    ws.append([])
    row += 1

    # Content rows
    for line in content.splitlines():
        line = line.strip()
        if not line:
            ws.append([])
            row += 1
            continue
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=3)
        cell = ws.cell(row=row, column=1, value=line)
        cell.font = body_font
        cell.alignment = right
        cell.border = cell_border
        ws.row_dimensions[row].height = 18
        row += 1

    ws.column_dimensions["A"].width = 24
    ws.column_dimensions["B"].width = 40
    ws.column_dimensions["C"].width = 24

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


# ═══════════════════════════════════════════════════════════
# Web search, scientific plots, template learning
# ═══════════════════════════════════════════════════════════
def web_search(query: str, max_results: int = 3) -> str:
    if not TAVILY_AVAILABLE or not TAVILY_API_KEY:
        return ""
    try:
        client = TavilyClient(api_key=TAVILY_API_KEY)
        response = client.search(query, max_results=max_results, include_answer=True)
        results = []
        if response.get("answer"):
            results.append(f"إجابة مختصرة: {response['answer']}")
        for res in response.get("results", []):
            results.append(f"- {res.get('title', '')}: {res.get('content', '')[:200]}")
        return "\n".join(results)
    except Exception as e:
        st.warning(f"فشل البحث: {e}")
        return ""

def _fig_to_bytes(fig) -> bytes:
    """Save a matplotlib figure to PNG bytes."""
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", dpi=130,
                facecolor=fig.get_facecolor())
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def _safe_eval_expr(expr_str: str, x_arr):
    """Safely evaluate a math expression with numpy vectors."""
    allowed = {
        'x': x_arr, 'abs': np.abs, 'sin': np.sin, 'cos': np.cos,
        'tan': np.tan, 'exp': np.exp, 'log': np.log, 'sqrt': np.sqrt,
        'pi': np.pi, 'e': np.e, 'arcsin': np.arcsin, 'arccos': np.arccos,
        'arctan': np.arctan, 'sinh': np.sinh, 'cosh': np.cosh,
        'tanh': np.tanh, 'log2': np.log2, 'log10': np.log10,
    }
    return eval(expr_str, {"__builtins__": {}}, allowed)


def generate_geometry_figure(shape: str, params: dict) -> bytes:
    """
    Draw a geometric shape with Matplotlib and return PNG bytes.
    Shapes: circle | rectangle | triangle | right_triangle |
            parallelogram | trapezoid
    """
    fig, ax = plt.subplots(figsize=(5, 5), facecolor='#f8fff9')
    ax.set_aspect('equal')
    ax.set_facecolor('#f8fff9')
    for spine in ax.spines.values():
        spine.set_color('#cccccc')
    ax.tick_params(colors='#555555')
    C_EDGE  = '#145a32'
    C_FILL  = '#d5f5e3'
    C_DIM1  = '#c0392b'
    C_DIM2  = '#1a5276'
    C_DIM3  = '#7d3c98'

    if shape == 'circle':
        r = float(params.get('r', 3))
        circ = plt.Circle((0, 0), r, edgecolor=C_EDGE, facecolor=C_FILL, lw=2.5)
        ax.add_patch(circ)
        ax.plot([0, r], [0, 0], color=C_DIM1, lw=1.8, ls='--')
        ax.text(r / 2, 0.22, f'r = {r}', color=C_DIM1, ha='center', fontsize=11, fontweight='bold')
        ax.set_xlim(-r*1.35, r*1.35); ax.set_ylim(-r*1.35, r*1.35)
        ax.set_title(f'دائرة  |  r = {r}  |  S = {3.14159*r**2:.3f}  |  P = {2*3.14159*r:.3f}', fontsize=11)

    elif shape == 'rectangle':
        w = float(params.get('w', 6)); h = float(params.get('h', 4))
        rect = mpatches.FancyBboxPatch((0, 0), w, h,
               boxstyle="square,pad=0", edgecolor=C_EDGE, facecolor=C_FILL, lw=2.5)
        ax.add_patch(rect)
        ax.annotate('', xy=(w, -0.3), xytext=(0, -0.3),
                    arrowprops=dict(arrowstyle='<->', color=C_DIM1, lw=1.5))
        ax.text(w/2, -0.55, f'L = {w}', ha='center', fontsize=11, color=C_DIM1, fontweight='bold')
        ax.annotate('', xy=(w+0.3, h), xytext=(w+0.3, 0),
                    arrowprops=dict(arrowstyle='<->', color=C_DIM2, lw=1.5))
        ax.text(w+0.6, h/2, f'l = {h}', ha='left', fontsize=11, color=C_DIM2, fontweight='bold', va='center')
        ax.set_xlim(-0.4, w+1); ax.set_ylim(-0.7, h+0.4)
        ax.set_title(f'مستطيل  |  S = {w*h}  |  P = {2*(w+h)}', fontsize=11)

    elif shape == 'triangle':
        a = float(params.get('a', 5)); h = float(params.get('h', 4)); b_side = float(params.get('b', 4.5))
        pts = np.array([[0,0],[a,0],[a/2,h],[0,0]])
        ax.fill(pts[:,0], pts[:,1], fc=C_FILL, ec=C_EDGE, lw=2.5)
        ax.plot([a/2, a/2], [0, h], color=C_DIM1, lw=1.5, ls='--')
        ax.text(a/2, -0.3, f'a = {a}', ha='center', fontsize=11, color=C_DIM1, fontweight='bold')
        ax.text(a/2+0.2, h/2, f'h = {h}', ha='left', fontsize=10, color=C_DIM2)
        ax.set_xlim(-0.5, a+0.5); ax.set_ylim(-0.5, h+0.4)
        ax.set_title(f'مثلث  |  S = {0.5*a*h:.3f}  |  a = {a}  h = {h}', fontsize=11)

    elif shape == 'right_triangle':
        a = float(params.get('a', 3)); b = float(params.get('b', 4))
        hyp = np.sqrt(a**2 + b**2)
        pts = np.array([[0,0],[a,0],[0,b],[0,0]])
        ax.fill(pts[:,0], pts[:,1], fc=C_FILL, ec=C_EDGE, lw=2.5)
        sq = 0.28
        ax.plot([sq, sq, 0], [0, sq, sq], color=C_EDGE, lw=1.5)
        ax.annotate('', xy=(a,-0.3), xytext=(0,-0.3),
                    arrowprops=dict(arrowstyle='<->', color=C_DIM1, lw=1.5))
        ax.text(a/2, -0.55, f'{a}', ha='center', fontsize=12, color=C_DIM1, fontweight='bold')
        ax.annotate('', xy=(-0.35, b), xytext=(-0.35, 0),
                    arrowprops=dict(arrowstyle='<->', color=C_DIM2, lw=1.5))
        ax.text(-0.6, b/2, f'{b}', ha='right', fontsize=12, color=C_DIM2, fontweight='bold', va='center')
        ax.text(a/2+0.2, b/2, f'c = {hyp:.3f}', fontsize=10, color=C_DIM3,
                rotation=-np.degrees(np.arctan2(b, a)))
        ax.set_xlim(-0.9, a+0.5); ax.set_ylim(-0.7, b+0.4)
        ax.set_title(f'مثلث قائم  |  وتر = {hyp:.3f}  |  S = {0.5*a*b:.3f}', fontsize=11)

    elif shape == 'parallelogram':
        b = float(params.get('b', 6)); h = float(params.get('h', 3)); sk = float(params.get('skew', 1.5))
        pts = np.array([[0,0],[b,0],[b+sk,h],[sk,h],[0,0]])
        ax.fill(pts[:,0], pts[:,1], fc=C_FILL, ec=C_EDGE, lw=2.5)
        ax.plot([sk, sk], [0, h], color=C_DIM1, lw=1.5, ls='--')
        ax.text(b/2+sk/2, -0.3, f'b = {b}', ha='center', fontsize=11, color=C_DIM1, fontweight='bold')
        ax.text(sk+0.15, h/2, f'h = {h}', ha='left', fontsize=11, color=C_DIM2, fontweight='bold', va='center')
        ax.set_xlim(-0.5, b+sk+0.5); ax.set_ylim(-0.5, h+0.4)
        ax.set_title(f'متوازي أضلاع  |  S = {b*h}  |  b = {b}  h = {h}', fontsize=11)

    elif shape == 'trapezoid':
        a = float(params.get('a', 6)); b = float(params.get('b', 3)); h = float(params.get('h', 3))
        off = (a - b) / 2
        pts = np.array([[0,0],[a,0],[a-off,h],[off,h],[0,0]])
        ax.fill(pts[:,0], pts[:,1], fc=C_FILL, ec=C_EDGE, lw=2.5)
        ax.plot([off, off], [0, h], color=C_DIM1, lw=1.5, ls='--')
        ax.text(a/2, -0.3, f'a = {a}', ha='center', fontsize=11, color=C_DIM1, fontweight='bold')
        ax.text(a/2, h+0.15, f'b = {b}', ha='center', fontsize=11, color=C_DIM2, fontweight='bold')
        ax.text(off-0.2, h/2, f'h = {h}', ha='right', fontsize=11, color=C_DIM3, fontweight='bold', va='center')
        ax.set_xlim(-0.5, a+0.5); ax.set_ylim(-0.5, h+0.5)
        ax.set_title(f'شبه منحرف  |  S = {0.5*(a+b)*h:.3f}  |  a={a}  b={b}  h={h}', fontsize=11)

    ax.grid(True, alpha=0.25, color='#aaaaaa', ls='--')
    return _fig_to_bytes(fig)


def generate_function_plot(expr_str: str, x_range=(-10, 10), label: str = "") -> bytes:
    """Plot a mathematical function f(x) and return PNG bytes."""
    x = np.linspace(x_range[0], x_range[1], 600)
    expr_clean = expr_str.replace('^', '**').replace('×', '*')
    try:
        y = _safe_eval_expr(expr_clean, x)
        y = np.where(np.abs(y) > 1e5, np.nan, y.astype(float))
    except Exception:
        return b""
    fig, ax = plt.subplots(figsize=(8, 4.5), facecolor='#f8fff9')
    ax.set_facecolor('#f8fff9')
    ax.axhline(0, color='#888', lw=0.9, zorder=1)
    ax.axvline(0, color='#888', lw=0.9, zorder=1)
    ax.plot(x, y, color='#145a32', lw=2.5, zorder=2,
            label=label or f'f(x) = {expr_str}')
    ax.grid(True, alpha=0.3, ls='--')
    ax.legend(fontsize=11)
    ax.set_title(f'f(x) = {expr_str}', fontsize=13)
    ax.set_xlabel('x', fontsize=11); ax.set_ylabel('f(x)', fontsize=11)
    return _fig_to_bytes(fig)


def auto_generate_plots(content: str, subject: str) -> list:
    """
    ENHANCED v5.1: Returns list of (type, data) tuples.
    type = 'plotly'  → data is a Plotly Figure
    type = 'image'   → data is PNG bytes (Matplotlib)
    Auto-detects functions and geometric shapes from generated text.
    """
    plots = []
    subject_lower = (subject or "").lower()
    is_math = any(k in subject_lower for k in [
        "رياضيات", "math", "mathematics", "فيزياء", "physics",
        "ميكانيك", "جبر", "هندسة", "algebra", "geometry"
    ])
    if not is_math:
        return plots

    # ── Plotly: explicit f(x)=... patterns ──
    func_pattern = r'f\s*\(\s*x\s*\)\s*=\s*([\d\.\w\s\^\+\-\*\/\(\)]+)'
    for expr in re.findall(r'f\s*\(\s*x\s*\)\s*=\s*([\d\.\w\s\^\+\-\*\/\(\)]+)', content)[:3]:
        expr_clean = expr.strip().replace('^', '**')
        try:
            x_v = np.linspace(-10, 10, 300)
            y_v = _safe_eval_expr(expr_clean, x_v)
            y_v = np.where(np.abs(y_v) > 1e5, np.nan, y_v.astype(float))
            fig = go.Figure(go.Scatter(
                x=x_v, y=y_v, mode='lines',
                name=f'f(x)={expr.strip()}',
                line=dict(color='#145a32', width=2.5)
            ))
            fig.update_layout(title=f'f(x) = {expr.strip()}',
                              xaxis_title='x', yaxis_title='f(x)',
                              template='plotly_white', height=360)
            plots.append(("plotly", fig))
        except Exception:
            pass

    # ── Statistical scatter from (x, y) pairs ──
    numbers = re.findall(r'(\d+)\s*,\s*(\d+)', content)
    if len(numbers) >= 3:
        df = pd.DataFrame(numbers[:20], columns=["X", "Y"]).astype(float)
        fig2 = px.scatter(df, x="X", y="Y",
                          title="تمثيل البيانات الإحصائية", template="plotly_white")
        fig2.update_traces(marker=dict(color='#c0392b', size=9))
        plots.append(("plotly", fig2))

    # ── Matplotlib geometry auto-detection ──
    geo_map = {
        "دائرة":           "circle",
        "circle":          "circle",
        "مثلث قائم":       "right_triangle",
        "right triangle":  "right_triangle",
        "مثلث":            "triangle",
        "triangle":        "triangle",
        "مستطيل":          "rectangle",
        "rectangle":       "rectangle",
        "متوازي أضلاع":    "parallelogram",
        "parallelogram":   "parallelogram",
        "شبه منحرف":       "trapezoid",
        "trapezoid":       "trapezoid",
    }
    content_lower = content.lower()
    drawn = set()
    for kw, shape in geo_map.items():
        if kw not in content_lower or shape in drawn:
            continue
        hit = content_lower.find(kw)
        nums = re.findall(r'(\d+(?:\.\d+)?)', content[hit:hit+100])
        nums = re.findall(r'(\d+(?:\.\d+)?)', content[hit:hit+100])
        p = {}
        if shape == 'circle' and nums:
            p['r'] = nums[0]
        elif shape in ('triangle', 'right_triangle') and len(nums) >= 2:
            p['a'] = nums[0]; p['b'] = nums[1]
            if len(nums) >= 3: p['h'] = nums[2]
        elif shape == 'rectangle' and len(nums) >= 2:
            p['w'] = nums[0]; p['h'] = nums[1]
        elif shape == 'trapezoid' and len(nums) >= 2:
            p['a'] = nums[0]; p['b'] = nums[1]
            if len(nums) >= 3: p['h'] = nums[2]
        elif shape == 'parallelogram' and len(nums) >= 2:
            p['b'] = nums[0]; p['h'] = nums[1]
        try:
            img = generate_geometry_figure(shape, p)
            if img:
                plots.append(("image", img))
                drawn.add(shape)
        except Exception:
            pass
    return plots


# Template learning functions
def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    if not _PDFPLUMBER_AVAILABLE:
        return ""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(pdf_bytes)
            tmp_path = tmp.name
        text = ""
        with pdfplumber.open(tmp_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
        os.unlink(tmp_path)
        return text
    except Exception as e:
        st.warning(f"PDF extraction error: {e}")
        return ""

def extract_text_from_image(image_bytes: bytes) -> str:
    if not _TESSERACT_AVAILABLE:
        return ""
    try:
        bio = io.BytesIO(image_bytes)
        im = Image.open(bio).convert("RGB")
        return pytesseract.image_to_string(im, lang="ara+eng+fra")
    except Exception as e:
        st.warning(f"OCR error: {e}")
        return ""

def analyze_template_structure(raw_text: str, template_type: str) -> dict:
    if not GROQ_API_KEY:
        return {"error": "Groq API missing"}
    prompt = f"""أنت خبير في تحليل هياكل الوثائق التعليمية الجزائرية.
النص المستخرج من القالب ({template_type}):
{raw_text[:3000]}

قم بتحليل الهيكل العام وأخرج بصيغة JSON:
{{
    "type": "lesson_plan" or "exam" or "exercise",
    "sections": ["قائمة", "العناوين", "الرئيسية"],
    "metadata": {{
        "has_table": true/false,
        "has_rtl": true,
        "has_equations": true/false
    }},
    "key_phrases": ["عبارات", "مهمة"],
    "suggested_prompt_template": "نموذج موجه يمكن استخدامه للتوليد وفق هذا القالب"
}}"""
    try:
        llm = get_llm(DEFAULT_GROQ_MODEL, GROQ_API_KEY)
        response = call_llm(llm, prompt)
        try:
            return json.loads(response)
        except:
            return {"type": "unknown", "sections": [], "metadata": {}, "key_phrases": [], "suggested_prompt_template": raw_text[:500]}
    except Exception as e:
        return {"error": str(e)}

def save_template(name: str, template_type: str, raw_text: str, structure: dict):
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS templates (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT,
            template_type TEXT,
            raw_text TEXT,
            structure_json TEXT,
            created_at TEXT
        )
    """)
    conn.execute(
        "INSERT INTO templates (name, template_type, raw_text, structure_json, created_at) VALUES (?,?,?,?,?)",
        (name, template_type, raw_text, json.dumps(structure, ensure_ascii=False), datetime.now().strftime("%Y-%m-%d %H:%M"))
    )
    conn.commit()
    conn.close()

def get_all_templates():
    conn = sqlite3.connect(DB_PATH)
    conn.execute("CREATE TABLE IF NOT EXISTS templates (id INTEGER PRIMARY KEY AUTOINCREMENT, name TEXT, template_type TEXT, raw_text TEXT, structure_json TEXT, created_at TEXT)")
    rows = conn.execute("SELECT id, name, template_type, created_at FROM templates ORDER BY created_at DESC").fetchall()
    conn.close()
    return rows

def get_template_by_id(tid: int):
    conn = sqlite3.connect(DB_PATH)
    row = conn.execute("SELECT raw_text, structure_json FROM templates WHERE id=?", (tid,)).fetchone()
    conn.close()
    if row:
        return row[0], json.loads(row[1])
    return None, None

# Audio helper
def audio_to_text(audio_bytes: bytes) -> str:
    if not GROQ_API_KEY:
        return ""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp:
            tmp.write(audio_bytes)
            tmp_path = tmp.name
        url = "https://api.groq.com/openai/v1/audio/transcriptions"
        headers = {"Authorization": f"Bearer {GROQ_API_KEY}"}
        with open(tmp_path, "rb") as f:
            files = {"file": f}
            data = {"model": "whisper-large-v3", "language": "auto"}
            response = requests.post(url, headers=headers, files=files, data=data)
        os.unlink(tmp_path)
        if response.status_code == 200:
            return response.json().get("text", "")
        return ""
    except Exception as e:
        st.error(f"خطأ في التعرف على الصوت: {e}")
        return ""

# Database
DB_PATH = "donia_smart.db"

def init_db():
    con = sqlite3.connect(DB_PATH)
    con.execute("""CREATE TABLE IF NOT EXISTS exercises (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        level TEXT, grade TEXT, branch TEXT, subject TEXT, lesson TEXT,
        ex_type TEXT, difficulty TEXT, content TEXT, created_at TEXT)""")
    con.execute("""CREATE TABLE IF NOT EXISTS lesson_plans (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        level TEXT, grade TEXT, subject TEXT, lesson TEXT,
        domain TEXT, duration TEXT, content TEXT, created_at TEXT)""")
    con.execute("""CREATE TABLE IF NOT EXISTS corrections (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        student_name TEXT, subject TEXT, grade_value REAL,
        total REAL, feedback TEXT, created_at TEXT)""")
    con.execute("""CREATE TABLE IF NOT EXISTS exams (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        level TEXT, grade TEXT, subject TEXT, semester TEXT,
        content TEXT, created_at TEXT)""")
    con.execute("""CREATE TABLE IF NOT EXISTS grade_books (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        class_name TEXT, subject TEXT, semester TEXT,
        data_json TEXT, created_at TEXT)""")
    con.execute("""CREATE TABLE IF NOT EXISTS templates (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT, template_type TEXT, raw_text TEXT,
        structure_json TEXT, created_at TEXT)""")
    con.execute("""CREATE TABLE IF NOT EXISTS rag_documents (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        content TEXT, content_type TEXT, metadata_json TEXT, created_at TEXT)""")
    con.commit()
    con.close()

def db_exec(sql, params=(), fetch=False):
    con = sqlite3.connect(DB_PATH)
    cur = con.execute(sql, params)
    con.commit()
    result = cur.fetchall() if fetch else None
    con.close()
    return result

def get_stats():
    total = (db_exec("SELECT COUNT(*) FROM exercises", fetch=True) or [(0,)])[0][0]
    plans = (db_exec("SELECT COUNT(*) FROM lesson_plans", fetch=True) or [(0,)])[0][0]
    exams = (db_exec("SELECT COUNT(*) FROM exams", fetch=True) or [(0,)])[0][0]
    corr = (db_exec("SELECT COUNT(*) FROM corrections", fetch=True) or [(0,)])[0][0]
    return total, plans, exams, corr

init_db()

# Helper functions (preserved from original)
def get_appreciation(grade, total=20):
    pct = grade / total * 100
    if pct >= 90:
        return "ممتاز"
    elif pct >= 75:
        return "جيد جداً"
    elif pct >= 65:
        return "جيد"
    elif pct >= 50:
        return "مقبول"
    else:
        return "ضعيف"

def calc_average(taqwim, fard, ikhtibhar):
    try:
        t = float(taqwim or 0)
        f = float(fard or 0)
        i = float(ikhtibhar or 0)
        return round((t * 1 + f * 1 + i * 2) / 4, 2)
    except (TypeError, ValueError):
        return 0.0

def safe_f(val, fmt=".2f") -> str:
    try:
        return format(float(val), fmt)
    except (TypeError, ValueError):
        return "—"

def render_with_latex(text):
    text = clean_latex(text)
    parts = re.split(r'(\$\$[\s\S]+?\$\$|\$[^\$\n]+?\$)', text)
    for part in parts:
        if part.startswith("$$") and part.endswith("$$"):
            st.latex(part[2:-2].strip())
        elif part.startswith("$") and part.endswith("$"):
            st.latex(part[1:-1].strip())
        elif part.strip():
            st.markdown(
                f'<div style="direction:rtl;text-align:right;'
                f'color:#111111;line-height:2;">{part}</div>',
                unsafe_allow_html=True)

def ocr_answer_sheet_image(image_bytes: bytes) -> str:
    if not _TESSERACT_AVAILABLE:
        return ""
    try:
        bio = io.BytesIO(image_bytes)
        im = Image.open(bio).convert("RGB")
        return pytesseract.image_to_string(im, lang="ara+eng+fra")
    except Exception:
        return ""

def build_class_stats(stus: list, cls_name: str) -> dict:
    avgs = [s['average'] for s in stus]
    passed = [a for a in avgs if a >= 10]
    dist = {"0-5": 0, "5-10": 0, "10-15": 0, "15-20": 0}
    for a in avgs:
        if a < 5:
            dist["0-5"] += 1
        elif a < 10:
            dist["5-10"] += 1
        elif a < 15:
            dist["10-15"] += 1
        else:
            dist["15-20"] += 1
    sorted_stus = sorted(stus, key=lambda x: x['average'], reverse=True)
    return {
        "name": cls_name,
        "total": len(stus),
        "avg": sum(avgs) / max(len(avgs), 1),
        "max": max(avgs) if avgs else 0.0,
        "min": min(avgs) if avgs else 0.0,
        "pass_rate": len(passed) / max(len(avgs), 1) * 100,
        "distribution": dist,
        "top5": [{"name": f"{s['nom']} {s['prenom']}", "avg": s['average']}
                 for s in sorted_stus[:5]],
        "students": stus,
    }

def parse_grade_book_excel(uploaded_file, sheet_name=None, merge_all_sheets=False) -> list:
    students = []
    try:
        uploaded_file.seek(0)
        wb = openpyxl.load_workbook(uploaded_file, data_only=True)
        if merge_all_sheets:
            for nm in wb.sheetnames:
                rows_list = list(wb[nm].iter_rows(values_only=True))
                part = _parse_rows_from_list(rows_list)
                for s in part:
                    s['sheet_source'] = nm
                students.extend(part)
            return students
        if sheet_name and sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
        else:
            ws = wb.active
        rows_list = list(ws.iter_rows(values_only=True))
        return _parse_rows_from_list(rows_list)
    except Exception:
        uploaded_file.seek(0)
        raw = uploaded_file.read()
        bio = io.BytesIO(raw)
        try:
            xl = pd.ExcelFile(bio)
        except Exception:
            try:
                df = pd.read_excel(bio, engine="openpyxl", header=None)
                rows_list = [tuple(row) for row in df.values]
                return _parse_rows_from_list(rows_list)
            except Exception:
                return []
        if merge_all_sheets:
            for nm in xl.sheet_names:
                df = pd.read_excel(xl, sheet_name=nm, header=None)
                rows_list = [tuple(row) for row in df.values]
                part = _parse_rows_from_list(rows_list)
                for s in part:
                    s['sheet_source'] = nm
                students.extend(part)
            return students
        sn = sheet_name if sheet_name and sheet_name in xl.sheet_names else xl.sheet_names[0]
        df = pd.read_excel(xl, sheet_name=sn, header=None)
        rows_list = [tuple(row) for row in df.values]
        return _parse_rows_from_list(rows_list)

def _parse_rows_from_list(rows_list) -> list:
    students = []
    data_started = False
    HEADER_MARKERS = {'matricule', 'رقم التعريف', 'اللقب', 'nom', 'prénom',
                      'الاسم', 'تقويم', 'فرض', 'اختبار', 'taqwim'}
    for i, row in enumerate(rows_list, 1):
        if not any(c is not None for c in row):
            continue
        row_strs = [str(c).strip().lower() if c is not None else '' for c in row]
        if not data_started:
            if any(m in row_strs for m in HEADER_MARKERS):
                data_started = True
            continue
        vals = list(row)
        if len(vals) < 4:
            continue
        nom = str(vals[1] or '').strip()
        if not nom or nom.lower() in ('اللقب', 'nom', 'prénom', 'الاسم'):
            continue
        try:
            stu = {
                'id': str(vals[0] or '').strip(),
                'nom': nom,
                'prenom': str(vals[2] or '').strip() if len(vals) > 2 else '',
                'dob': str(vals[3] or '').strip() if len(vals) > 3 else '',
                'taqwim': float(vals[4]) if len(vals) > 4 and vals[4] is not None else 0.0,
                'fard': float(vals[5]) if len(vals) > 5 and vals[5] is not None else 0.0,
                'ikhtibhar': float(vals[6]) if len(vals) > 6 and vals[6] is not None else 0.0,
            }
            stu['average'] = calc_average(stu['taqwim'], stu['fard'], stu['ikhtibhar'])
            stu['apprec'] = get_appreciation(stu['average'])
            students.append(stu)
        except (ValueError, TypeError):
            continue
    return students

def list_excel_sheet_names(uploaded_file) -> list:
    try:
        uploaded_file.seek(0)
        wb = openpyxl.load_workbook(uploaded_file, read_only=True, data_only=True)
        names = list(wb.sheetnames)
        wb.close()
        return names
    except Exception:
        uploaded_file.seek(0)
        raw = uploaded_file.read()
        bio = io.BytesIO(raw)
        try:
            xl = pd.ExcelFile(bio)
            return list(xl.sheet_names)
        except Exception:
            return []

def llm_output_language_clause(subject: str) -> str:
    rtl, lang = get_pdf_mode_for_subject(subject)
    if rtl:
        return "قاعدة إلزامية: اكتب كل المحتوى (العناوين، الأسئلة، الشروح) بالعربية الفصحى الواضحة."
    else:
        return f"Mandatory: produce the ENTIRE output in {lang}. Do not use Arabic."

# Curriculum data
CURRICULUM = {
    "الطور الابتدائي": {
        "grades": ["السنة الأولى", "السنة الثانية", "السنة الثالثة",
                   "السنة الرابعة", "السنة الخامسة"],
        "subjects": {
            "السنة الأولى": ["اللغة العربية", "الرياضيات", "التربية الإسلامية",
                             "التربية المدنية", "التربية التشكيلية", "التربية البدنية"],
            "السنة الثانية": ["اللغة العربية", "الرياضيات", "التربية الإسلامية",
                              "التربية المدنية", "التربية التشكيلية", "التربية البدنية"],
            "السنة الثالثة": ["اللغة العربية", "الرياضيات", "التربية الإسلامية",
                              "التربية المدنية", "اللغة الفرنسية",
                              "التربية العلمية والتكنولوجية", "التاريخ والجغرافيا"],
            "السنة الرابعة": ["اللغة العربية", "الرياضيات", "التربية الإسلامية",
                              "التربية المدنية", "اللغة الفرنسية",
                              "التربية العلمية والتكنولوجية", "التاريخ والجغرافيا"],
            "السنة الخامسة": ["اللغة العربية", "الرياضيات", "التربية الإسلامية",
                              "التربية المدنية", "اللغة الفرنسية",
                              "التربية العلمية والتكنولوجية", "التاريخ والجغرافيا"],
        },
        "branches": None,
    },
    "الطور المتوسط": {
        "grades": ["السنة الأولى متوسط", "السنة الثانية متوسط",
                   "السنة الثالثة متوسط", "السنة الرابعة متوسط (شهادة)"],
        "subjects": {
            "_default": ["اللغة العربية وآدابها", "الرياضيات",
                         "العلوم الفيزيائية والتكنولوجية", "العلوم الطبيعية والحياة",
                         "التاريخ والجغرافيا", "الاجتماعيات",
                         "التربية الإسلامية", "التربية المدنية",
                         "اللغة الفرنسية", "اللغة الإنجليزية",
                         "التربية التشكيلية", "التربية الموسيقية", "الإعلام الآلي"]
        },
        "branches": None,
    },
    "الطور الثانوي": {
        "grades": ["السنة الأولى ثانوي (جذع مشترك)",
                   "السنة الثانية ثانوي", "السنة الثالثة ثانوي (بكالوريا)"],
        "subjects": None,
        "branches": {
            "السنة الأولى ثانوي (جذع مشترك)": {
                "جذع مشترك علوم وتكنولوجيا": [
                    "الرياضيات", "العلوم الفيزيائية", "العلوم الطبيعية والحياة",
                    "اللغة العربية", "اللغة الفرنسية", "اللغة الإنجليزية",
                    "التاريخ والجغرافيا", "التربية الإسلامية", "الإعلام الآلي"],
                "جذع مشترك آداب وفلسفة": [
                    "اللغة العربية وآدابها", "الفلسفة", "التاريخ والجغرافيا",
                    "اللغة الفرنسية", "اللغة الإنجليزية",
                    "التربية الإسلامية", "الرياضيات"],
            },
            "السنة الثانية ثانوي": {
                "شعبة علوم تجريبية": ["الرياضيات", "العلوم الفيزيائية",
                    "العلوم الطبيعية والحياة", "اللغة العربية", "اللغة الفرنسية",
                    "اللغة الإنجليزية", "التاريخ والجغرافيا", "التربية الإسلامية"],
                "شعبة رياضيات": ["الرياضيات", "العلوم الفيزيائية",
                    "العلوم الطبيعية والحياة", "اللغة العربية",
                    "اللغة الفرنسية", "اللغة الإنجليزية"],
                "شعبة تقني رياضي": ["الرياضيات", "العلوم الفيزيائية", "التكنولوجيا",
                    "اللغة العربية", "اللغة الفرنسية", "اللغة الإنجليزية"],
                "شعبة آداب وفلسفة": ["اللغة العربية وآدابها", "الفلسفة",
                    "التاريخ والجغرافيا", "علم الاجتماع والنفس",
                    "اللغة الفرنسية", "اللغة الإنجليزية", "التربية الإسلامية"],
                "شعبة لغات أجنبية": ["اللغة الفرنسية", "اللغة الإنجليزية",
                    "اللغة الإسبانية", "اللغة الألمانية", "اللغة الإيطالية",
                    "اللغة العربية", "التاريخ والجغرافيا", "الفلسفة"],
                "شعبة تسيير واقتصاد": ["الاقتصاد والمناجمنت", "المحاسبة والمالية",
                    "الرياضيات", "القانون", "اللغة العربية",
                    "اللغة الفرنسية", "اللغة الإنجليزية"],
            },
            "السنة الثالثة ثانوي (بكالوريا)": {
                "شعبة علوم تجريبية": ["الرياضيات", "العلوم الفيزيائية",
                    "العلوم الطبيعية والحياة", "اللغة العربية", "اللغة الفرنسية",
                    "اللغة الإنجليزية", "التاريخ والجغرافيا", "التربية الإسلامية"],
                "شعبة رياضيات": ["الرياضيات", "العلوم الفيزيائية",
                    "العلوم الطبيعية والحياة", "اللغة العربية",
                    "اللغة الفرنسية", "اللغة الإنجليزية"],
                "شعبة تقني رياضي": ["الرياضيات", "العلوم الفيزيائية", "التكنولوجيا",
                    "اللغة العربية", "اللغة الفرنسية", "اللغة الإنجليزية"],
                "شعبة آداب وفلسفة": ["اللغة العربية وآدابها", "الفلسفة",
                    "التاريخ والجغرافيا", "علم الاجتماع والنفس",
                    "اللغة الفرنسية", "اللغة الإنجليزية", "التربية الإسلامية"],
                "شعبة لغات أجنبية": ["اللغة الفرنسية", "اللغة الإنجليزية",
                    "اللغة الإسبانية", "اللغة الألمانية", "اللغة الإيطالية",
                    "اللغة العربية", "التاريخ والجغرافيا", "الفلسفة"],
                "شعبة تسيير واقتصاد": ["الاقتصاد والمناجمنت", "المحاسبة والمالية",
                    "الرياضيات", "القانون", "اللغة العربية",
                    "اللغة الفرنسية", "اللغة الإنجليزية"],
            },
        },
    },
}

DOMAINS = {
    "الرياضيات": ["أنشطة عددية", "أنشطة جبرية", "أنشطة هندسية", "أنشطة إحصائية"],
    "العلوم الفيزيائية والتكنولوجية": ["المادة", "الكهرباء", "الضوء", "الميكانيك"],
    "العلوم الطبيعية والحياة": ["الوحدة والتنوع", "التغذية والهضم", "التوليد", "البيئة"],
    "اللغة العربية وآدابها": ["فهم المكتوب", "الإنتاج الكتابي", "الظاهرة اللغوية", "الميدان الأدبي"],
}

# Excel generators
def generate_grade_book_excel(students: list, class_name: str, subject: str, semester: str, school_name: str) -> bytes:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = class_name[:31]
    title_font = Font(name="Arial", bold=True, size=11)
    header_font = Font(name="Arial", bold=True, size=10, color="FFFFFF")
    body_font = Font(name="Arial", size=10)
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    right = Alignment(horizontal="right", vertical="center")
    thin = Side(style="thin", color="000000")
    border = Border(top=thin, bottom=thin, left=thin, right=thin)
    purple_fill = PatternFill("solid", fgColor="764ba2")
    light_fill = PatternFill("solid", fgColor="f0f0ff")
    ws.merge_cells("A1:I1")
    ws["A1"] = "الجمهورية الجزائرية الديمقراطية الشعبية"
    ws["A1"].font = title_font
    ws["A1"].alignment = center
    ws["A1"].fill = light_fill
    ws.merge_cells("A2:I2")
    ws["A2"] = "وزارة التربية الوطنية"
    ws["A2"].font = title_font
    ws["A2"].alignment = center
    ws.merge_cells("A3:I3")
    ws["A3"] = f"المؤسسة: {school_name}"
    ws["A3"].font = title_font
    ws["A3"].alignment = center
    ws.merge_cells("A4:I4")
    ws["A4"] = f"دفتر التنقيط | القسم: {class_name} | المادة: {subject} | {semester}"
    ws["A4"].font = title_font
    ws["A4"].alignment = center
    ws["A4"].fill = PatternFill("solid", fgColor="e8e8ff")
    ws.append([])
    headers = ["الرقم", "اللقب", "الاسم", "تاريخ الميلاد",
               "تقويم /20", "فرض /20", "اختبار /20", "المعدل /20", "التقديرات"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=6, column=col, value=h)
        cell.font = header_font
        cell.alignment = center
        cell.fill = purple_fill
        cell.border = border
    ws.row_dimensions[6].height = 30
    for idx, stu in enumerate(students, 1):
        row = 6 + idx
        avg = stu.get('average', 0)
        apprec = get_appreciation(avg)
        values = [
            idx,
            stu.get('nom', ''),
            stu.get('prenom', ''),
            str(stu.get('dob', '')),
            stu.get('taqwim', ''),
            stu.get('fard', ''),
            stu.get('ikhtibhar', ''),
            avg,
            apprec,
        ]
        for col, val in enumerate(values, 1):
            cell = ws.cell(row=row, column=col, value=val)
            cell.font = body_font
            cell.border = border
            cell.alignment = center if col not in [2, 3] else right
            if idx % 2 == 0:
                cell.fill = PatternFill("solid", fgColor="f8f8ff")
        ws.row_dimensions[row].height = 22
    last_data = 6 + len(students)
    stat_row = last_data + 2
    avgs_all = [s.get('average', 0) for s in students]
    stats = [
        ("عدد التلاميذ", len(students)),
        ("معدل القسم", round(sum(avgs_all) / max(len(avgs_all), 1), 2)),
        ("الناجحون", sum(1 for a in avgs_all if a >= 10)),
    ]
    for i, (label, val) in enumerate(stats):
        lc = ws.cell(row=stat_row + i, column=1, value=label)
        vc = ws.cell(row=stat_row + i, column=2, value=val)
        lc.font = Font(bold=True, name="Arial", size=10)
        vc.font = Font(bold=True, name="Arial", size=10, color="764ba2")
        lc.fill = light_fill
        vc.fill = light_fill
        lc.border = border
        vc.border = border
    widths = [8, 16, 16, 14, 10, 10, 10, 10, 12]
    for col, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(col)].width = w
    ws.sheet_view.rightToLeft = True
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()

def generate_multi_sheet_grade_book(classes_data: list, school_name: str, subject: str, semester: str) -> bytes:
    wb = openpyxl.Workbook()
    wb.remove(wb.active)
    for cls_data in classes_data:
        students = cls_data.get('students', [])
        class_name = cls_data.get('name', 'قسم')
        sheet_name = class_name[:31]
        ws = wb.create_sheet(title=sheet_name)
        title_font = Font(name="Arial", bold=True, size=11)
        header_font = Font(name="Arial", bold=True, size=10, color="FFFFFF")
        body_font = Font(name="Arial", size=10)
        center = Alignment(horizontal="center", vertical="center", wrap_text=True)
        right = Alignment(horizontal="right", vertical="center")
        thin = Side(style="thin", color="000000")
        border = Border(top=thin, bottom=thin, left=thin, right=thin)
        purple_fill = PatternFill("solid", fgColor="764ba2")
        light_fill = PatternFill("solid", fgColor="f0f0ff")
        ws.merge_cells("A1:I1")
        ws["A1"] = "الجمهورية الجزائرية الديمقراطية الشعبية"
        ws["A1"].font = title_font
        ws["A1"].alignment = center
        ws["A1"].fill = light_fill
        ws.merge_cells("A2:I2")
        ws["A2"] = "وزارة التربية الوطنية"
        ws["A2"].font = title_font
        ws["A2"].alignment = center
        ws.merge_cells("A3:I3")
        ws["A3"] = f"المؤسسة: {school_name}"
        ws["A3"].font = title_font
        ws["A3"].alignment = center
        ws.merge_cells("A4:I4")
        ws["A4"] = f"دفتر التنقيط | القسم: {class_name} | المادة: {subject} | {semester}"
        ws["A4"].font = title_font
        ws["A4"].alignment = center
        ws["A4"].fill = PatternFill("solid", fgColor="e8e8ff")
        ws.append([])
        headers = ["الرقم", "اللقب", "الاسم", "تاريخ الميلاد",
                   "تقويم /20", "فرض /20", "اختبار /20", "المعدل /20", "التقديرات"]
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=6, column=col, value=h)
            cell.font = header_font
            cell.alignment = center
            cell.fill = purple_fill
            cell.border = border
        ws.row_dimensions[6].height = 30
        for idx, stu in enumerate(students, 1):
            row = 6 + idx
            avg = stu.get('average', 0)
            apprec = get_appreciation(avg)
            values = [
                idx,
                stu.get('nom', ''),
                stu.get('prenom', ''),
                str(stu.get('dob', '')),
                stu.get('taqwim', ''),
                stu.get('fard', ''),
                stu.get('ikhtibhar', ''),
                avg,
                apprec,
            ]
            for col, val in enumerate(values, 1):
                cell = ws.cell(row=row, column=col, value=val)
                cell.font = body_font
                cell.border = border
                cell.alignment = center if col not in [2, 3] else right
                if idx % 2 == 0:
                    cell.fill = PatternFill("solid", fgColor="f8f8ff")
            ws.row_dimensions[row].height = 22
        last_data = 6 + len(students)
        stat_row = last_data + 2
        avgs_all = [s.get('average', 0) for s in students]
        stats = [
            ("عدد التلاميذ", len(students)),
            ("معدل القسم", round(sum(avgs_all) / max(len(avgs_all), 1), 2)),
            ("الناجحون", sum(1 for a in avgs_all if a >= 10)),
        ]
        for i, (label, val) in enumerate(stats):
            lc = ws.cell(row=stat_row + i, column=1, value=label)
            vc = ws.cell(row=stat_row + i, column=2, value=val)
            lc.font = Font(bold=True, name="Arial", size=10)
            vc.font = Font(bold=True, name="Arial", size=10, color="764ba2")
            lc.fill = light_fill
            vc.fill = light_fill
            lc.border = border
            vc.border = border
        widths = [8, 16, 16, 14, 10, 10, 10, 10, 12]
        for col, w in enumerate(widths, 1):
            ws.column_dimensions[get_column_letter(col)].width = w
        ws.sheet_view.rightToLeft = True
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()

# Floating assistant
def render_floating_assistant():
    if "assistant_messages" not in st.session_state:
        st.session_state.assistant_messages = [
            {"role": "assistant", "content": "🌟 مرحباً بك في DONIA MIND v5.0! أنا مساعدك الذكي. يمكنني مساعدتك في إعداد المذكرات، توليد الاختبارات، وتحليل النتائج."}
        ]
    if "assistant_open" not in st.session_state:
        st.session_state.assistant_open = False
    button_html = """
    <div class="floating-assistant" id="assistantToggle" onclick="document.getElementById('assistantChat').style.display = document.getElementById('assistantChat').style.display === 'none' ? 'block' : 'none';">
        <div class="assistant-bubble">
            <svg viewBox="0 0 80 80" xmlns="http://www.w3.org/2000/svg">
                <rect x="15" y="18" width="50" height="44" rx="14" fill="#ffffff" stroke="#c0392b" stroke-width="2"/>
                <circle cx="31" cy="36" r="5" fill="#145a32"/>
                <circle cx="49" cy="36" r="5" fill="#145a32"/>
                <circle cx="32" cy="35" r="2" fill="white"/>
                <circle cx="50" cy="35" r="2" fill="white"/>
                <path d="M30 52 Q40 58 50 52" stroke="#c0392b" stroke-width="2.5" fill="none"/>
                <line x1="40" y1="18" x2="40" y2="12" stroke="#c0392b" stroke-width="2"/>
                <circle cx="40" cy="10" r="3" fill="#c0392b"/>
            </svg>
        </div>
    </div>
    """
    st.markdown(button_html, unsafe_allow_html=True)
    with st.container():
        st.markdown('<div id="assistantChat" style="display: none;">', unsafe_allow_html=True)
        with st.chat_message("assistant", avatar="🤖"):
            st.markdown("🌟 مرحباً بك في DONIA MIND v5.0! أنا مساعدك الذكي.")
            st.markdown("يمكنني مساعدتك في:")
            st.markdown("- 📝 إعداد المذكرات")
            st.markdown("- 📄 توليد الاختبارات")
            st.markdown("- 📊 تحليل النتائج")
            st.markdown("- ✅ تصحيح الإجابات")
            st.markdown("- 🎤 يمكنك أيضاً استخدام الإدخال الصوتي (زر الميكروفون في الشريط الجانبي)")
        user_input = st.chat_input("اكتب سؤالك هنا...", key="assistant_input")
        if user_input:
            st.session_state.assistant_messages.append({"role": "user", "content": user_input})
            with st.chat_message("user"):
                st.markdown(user_input)
            with st.chat_message("assistant"):
                response = generate_assistant_response(user_input)
                st.markdown(response)
                st.session_state.assistant_messages.append({"role": "assistant", "content": response})
        st.markdown('</div>', unsafe_allow_html=True)

def generate_assistant_response(query: str) -> str:
    if not GROQ_API_KEY:
        return "⚠️ عذراً، مفتاح API غير متوفر."
    try:
        prompt = f"""أنت مساعد تربوي ذكي متخصص في المنظومة التعليمية الجزائرية.
        المستخدم يسأل: {query}
        قدم إجابة مفيدة ودقيقة حول المذكرات، الاختبارات، طرق التدريس، المنهاج الجزائري، التنقيط والتقييم.
        كن مختصراً وواضحاً."""
        llm = get_llm(DEFAULT_GROQ_MODEL, GROQ_API_KEY)
        response = call_llm(llm, prompt)
        return response
    except Exception as e:
        return f"❌ حدث خطأ: {str(e)}"

# ========== PAGE CONFIG & CSS ==========
st.set_page_config(page_title="DONIA MIND — المعلم الذكي v5.0", page_icon="🎓",
                   layout="wide", initial_sidebar_state="expanded")

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Amiri:wght@400;700&family=Cairo:wght@400;600;700;800&family=Tajawal:wght@400;500;700;800&family=Montserrat:wght@400;600;700;800;900&display=swap');

#MainMenu{visibility:hidden!important}
footer{visibility:hidden!important}
header{visibility:hidden!important}
.stDeployButton{display:none!important}
[data-testid="stToolbar"]{display:none!important}
[data-testid="stDecoration"]{display:none!important}
[data-testid="stStatusWidget"]{display:none!important}
a[href*="streamlit.io"]{display:none!important}

*,*::before,*::after{font-family:'Cairo','Amiri','Tajawal',sans-serif!important}
.stApp{background:#ffffff;color:#111111;}
.main{direction:rtl;text-align:right;color:#111111!important}
.block-container{color:#111111!important;background:#ffffff;}

@media (max-width: 768px) {
    .stColumn { width: 100% !important; flex: 1 1 100% !important; margin-bottom: 1rem; }
    .stButton > button { width: 100% !important; padding: 0.8rem !important; font-size: 0.9rem !important; }
    .stTextInput > div > div > input, .stTextArea > div > div > textarea { font-size: 16px !important; }
    .stSelectbox > div > div { font-size: 16px !important; }
    .title-card h1 { font-size: 1.5rem !important; }
    .donia-slogan-ar { font-size: 1.1rem !important; }
    .floating-assistant { bottom: 20px; right: 10px; }
    .assistant-bubble { width: 50px; height: 50px; }
    .assistant-bubble svg { width: 32px; height: 32px; }
}

.stMarkdown, .stTextInput, .stTextArea, .stSelectbox, .stRadio, .stCheckbox, .stButton, .stDataFrame {
    direction: rtl;
    text-align: right;
}
.ltr-text {
    direction: ltr;
    text-align: left;
}

h1{color:#c0392b!important;font-weight:800!important}
h2{color:#145a32!important;font-weight:700!important}
h3{color:#1e8449!important;font-weight:700!important}

.title-card{
  background:linear-gradient(135deg,#145a32 0%,#1e8449 50%,#27ae60 100%);
  padding:1.75rem 2rem;border-radius:24px;text-align:center;
  margin-bottom:1rem;box-shadow:0 16px 48px rgba(20,90,50,.45);
  border:3px solid #c0392b;
}
.title-card h1{color:#ffffff!important;font-size:2.05rem;font-weight:800;margin:0;letter-spacing:.02em}
.title-card p{color:rgba(255,255,255,.92);font-size:.96rem;margin:.45rem 0 0;line-height:1.65}

.welcome-banner{
  background:linear-gradient(135deg,#fdfefe,#f9f9f9);
  border:2px solid #27ae60;border-left:8px solid #c0392b;
  border-radius:14px;padding:1.1rem 1.5rem;margin:.75rem 0 1.25rem;
  direction:rtl;text-align:right;
  font-size:1.05rem;font-weight:600;color:#145a32;
  box-shadow:0 4px 16px rgba(20,90,50,.12);
}

.floating-assistant {
  position: fixed;
  bottom: 80px;
  right: 20px;
  z-index: 1000;
  cursor: pointer;
  transition: all 0.3s ease;
}
.floating-assistant:hover { transform: scale(1.05); }
.assistant-bubble {
  background: linear-gradient(135deg, #145a32, #1e8449);
  border-radius: 50%;
  width: 60px;
  height: 60px;
  display: flex;
  align-items: center;
  justify-content: center;
  box-shadow: 0 4px 20px rgba(20,90,50,.4);
  border: 2px solid #c0392b;
  animation: pulse 2s ease-in-out infinite;
}
@keyframes pulse {
  0%,100%{box-shadow:0 4px 20px rgba(39,174,96,.4)}
  50%{box-shadow:0 8px 30px rgba(192,57,43,.6)}
}
.assistant-bubble svg { width: 40px; height: 40px; }

.donia-robot-wrap{display:flex;justify-content:center;align-items:center;margin:.75rem 0}
.donia-robot{
  width:88px;height:88px;border-radius:22px;
  background:linear-gradient(180deg,#145a32,#1e8449);
  box-shadow:0 0 28px rgba(39,174,96,.55), inset 0 1px 0 rgba(255,255,255,.12);
  display:flex;align-items:center;justify-content:center;
  animation:doniaPulse 2.2s ease-in-out infinite;
  border:2px solid rgba(192,57,43,.6);
}
.donia-robot svg{width:64px;height:64px;opacity:.95}
@keyframes doniaPulse{
  0%,100%{transform:scale(1);box-shadow:0 0 28px rgba(39,174,96,.45)}
  50%{transform:scale(1.04);box-shadow:0 0 44px rgba(39,174,96,.85)}
}

div.stButton>button{
  background:linear-gradient(135deg,#1e8449,#145a32)!important;color:#ffffff!important;
  border:none!important;border-radius:18px!important;
  padding:0.85rem 1.65rem!important;min-height:3.1rem!important;
  font-weight:800!important;font-size:1.02rem!important;width:100%!important;
  transition:transform .22s, box-shadow .22s!important;
  box-shadow:0 6px 22px rgba(30,132,73,.45)!important;
}
div.stButton>button:hover{
  transform:translateY(-3px)!important;
  box-shadow:0 12px 36px rgba(192,57,43,.5)!important;
  background:linear-gradient(135deg,#c0392b,#922b21)!important;
}

.stat-card{background:linear-gradient(135deg,rgba(30,132,73,.1),rgba(39,174,96,.08));
  border:2px solid #27ae60;border-radius:16px;
  padding:1.1rem;text-align:center;margin-bottom:.75rem}
.stat-card h2{font-size:1.85rem;margin:0;color:#145a32!important}
.stat-card p{margin:0;color:#333;font-size:.86rem}

.feature-card{background:#f9f9f9;border:1px solid #27ae60;
  border-right:5px solid #1e8449;
  border-radius:16px;padding:1.25rem;margin:.55rem 0;
  direction:rtl;text-align:right;color:#111}
.feature-card h4{color:#1e8449;margin:0 0 .45rem;font-size:1.02rem}

.result-box{background:#f9f9f9;border:1px solid rgba(30,132,73,.3);
  border-radius:16px;padding:1.45rem;direction:rtl;text-align:right;
  color:#111;line-height:2;margin:.85rem 0}

.db-item{background:#f4f9f4;border-right:4px solid #1e8449;
  border-radius:10px;padding:.85rem 1.05rem;margin:.45rem 0;
  direction:rtl;text-align:right;color:#111}

.error-box{background:rgba(192,57,43,.08);border:2px solid #c0392b;
  border-radius:12px;padding:1rem;direction:rtl;text-align:right;
  color:#922b21;margin:.65rem 0;font-weight:600}
.success-box{background:rgba(30,132,73,.08);border:2px solid #27ae60;
  border-radius:12px;padding:1rem;direction:rtl;text-align:right;
  color:#145a32;margin:.65rem 0;font-weight:600}
.warn-box{background:rgba(243,156,18,.1);border:2px solid #f39c12;
  border-radius:12px;padding:1rem;direction:rtl;text-align:right;
  color:#784212;margin:.65rem 0}
.template-box{background:rgba(30,132,73,.06);border:2px dashed #27ae60;
  border-radius:14px;padding:1.05rem;direction:rtl;text-align:right;
  color:#145a32;margin:.65rem 0;font-size:.9rem;line-height:1.85}

.grade-A{color:#1e8449;font-weight:700}
.grade-B{color:#2e86c1;font-weight:700}
.grade-C{color:#d4ac0d;font-weight:700}
.grade-D{color:#c0392b;font-weight:700}

section[data-testid="stSidebar"]{
  direction:rtl;
  background:linear-gradient(180deg,#f4fbf6,#eaf6ee)!important;
  border-left:4px solid #27ae60;
}
@media (max-width: 768px) {
    section[data-testid="stSidebar"] { width: 85% !important; }
}
section[data-testid="stSidebar"] .stMarkdown{text-align:right;color:#145a32}

.stTabs [data-baseweb="tab"]{direction:rtl;font-size:.9rem;font-weight:700;color:#145a32}
.stTabs [data-baseweb="tab"][aria-selected="true"]{
  border-bottom:3px solid #c0392b!important;color:#c0392b!important}

.stSelectbox label,.stTextInput label,.stTextArea label,
.stNumberInput label,.stSlider label,.stFileUploader label,.stRadio label{
  direction:rtl;text-align:right;color:#145a32!important;font-weight:700}

.api-book-widget{
  background:linear-gradient(135deg,#f4fbf6,#eaf6ee);
  border:2px solid #27ae60;border-radius:16px;
  padding:1.1rem 1.2rem;text-align:center;margin:.5rem 0;
}
.api-book-icon{font-size:2.4rem;display:block;margin-bottom:.35rem}
.api-book-slogan{font-size:1rem;font-weight:800;color:#145a32;
  display:block;letter-spacing:.03em}
.api-book-status-active{
  display:block;margin-top:.4rem;font-size:.88rem;font-weight:700;
  color:#1e8449;background:#d5f5e3;border-radius:8px;padding:.2rem .7rem;
}
.api-book-status-inactive{
  display:block;margin-top:.4rem;font-size:.88rem;font-weight:700;
  color:#c0392b;background:#fdecea;border-radius:8px;padding:.2rem .7rem;
}

.donia-social{display:flex;flex-wrap:wrap;gap:.45rem;justify-content:center;margin:.35rem 0}
.donia-social a{
  display:inline-block;padding:.35rem .75rem;border-radius:12px;
  background:#145a32;color:#ffffff!important;font-weight:700;font-size:.82rem;
  text-decoration:none!important;border:1px solid #27ae60;
  transition:transform .2s,box-shadow .2s;
}
.donia-social a:hover{
  transform:translateY(-2px);
  box-shadow:0 6px 18px rgba(192,57,43,.4);
  background:#c0392b!important;
}

.donia-ip-footer{
  text-align:center;font-size:.85rem;color:#145a32;font-weight:600;
  padding:1.2rem 0 .5rem;margin-top:1.5rem;
  border-top:3px solid #27ae60;
  background:linear-gradient(90deg,#f4fbf6,#fef9f9,#f4fbf6);
  border-radius:0 0 12px 12px;
}
.donia-footer-social{display:flex;flex-wrap:wrap;gap:.6rem;justify-content:center;margin:.5rem 0}
.donia-footer-social a{
  display:inline-flex;align-items:center;gap:.3rem;
  padding:.4rem .9rem;border-radius:20px;
  background:#145a32;color:#ffffff!important;font-weight:700;font-size:.82rem;
  text-decoration:none!important;transition:background .2s,transform .2s;
}
.donia-footer-social a:hover{background:#c0392b!important;transform:translateY(-2px)}

.donia-slogan-bar{
  display:flex;flex-direction:column;align-items:center;
  gap:.3rem;padding:.9rem 1.5rem;margin:.6rem 0;
  background:linear-gradient(90deg,#145a32 0%,#1e8449 45%,#c0392b 100%);
  border-radius:14px;
  box-shadow:0 4px 20px rgba(20,90,50,.3);
}
.donia-slogan-ar{
  font-family:'Cairo','Amiri',sans-serif;
  font-size:1.35rem;font-weight:800;
  color:#ffffff;letter-spacing:.04em;
  text-shadow:0 2px 6px rgba(0,0,0,.3);
}
.donia-slogan-en{
  font-family:'Montserrat',sans-serif;
  font-size:.9rem;font-weight:600;
  color:rgba(255,255,255,.88);letter-spacing:.18em;
  text-transform:uppercase;
}
.donia-slogan-divider{
  width:40px;height:2px;
  background:rgba(255,255,255,.55);border-radius:2px;
}
</style>
""", unsafe_allow_html=True)

# ========== SIDEBAR ==========
with st.sidebar:
    _logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets", "logo_donia.jpg")
    if os.path.isfile(_logo_path):
        st.image(_logo_path, width=220, caption="DONIA LABS TECH v5.0")

    try:
        import qrcode
        from io import BytesIO
        qr = qrcode.QRCode(version=1, box_size=4, border=2)
        qr.add_data(APP_URL)
        qr.make(fit=True)
        qr_img = qr.make_image(fill_color="#145a32", back_color="white")
        qr_buf = BytesIO()
        qr_img.save(qr_buf, format="PNG")
        qr_buf.seek(0)
        st.image(qr_buf, caption="مسح للوصول السريع", width=120)
    except Exception:
        st.caption("📱 مسح للوصول للتطبيق")

    st.markdown("## ⚙️ الإعدادات العامة")

    st.markdown("### 🔌 حالة الاتصال")
    col1, col2 = st.columns(2)
    with col1:
        if GROQ_API_KEY:
            st.markdown('<div class="success-box" style="text-align:center">✅ Groq: متصل</div>', unsafe_allow_html=True)
        else:
            st.markdown('<div class="error-box" style="text-align:center">❌ Groq: غير متصل</div>', unsafe_allow_html=True)
    with col2:
        arcee_connected = test_arcee_connection()
        if arcee_connected:
            st.markdown('<div class="success-box" style="text-align:center">✅ Arcee: متصل</div>', unsafe_allow_html=True)
        else:
            st.markdown('<div class="error-box" style="text-align:center">❌ Arcee: غير متصل</div>', unsafe_allow_html=True)

    st.markdown("### 🎤 إدخال صوتي")
    if MIC_AVAILABLE:
        _mic_result = mic_recorder(start_prompt="🎙️ اضغط للتسجيل", stop_prompt="⏹️ إيقاف", key="mic_recorder")
        audio_bytes = _mic_result.get("bytes") if isinstance(_mic_result, dict) else _mic_result
        if audio_bytes:
            with st.spinner("جاري تحويل الصوت إلى نص..."):
                transcribed = audio_to_text(audio_bytes)
                if transcribed:
                    st.success(f"تم التعرف: {transcribed[:100]}...")
                    st.session_state["voice_text"] = transcribed
                    st.info("يمكنك استخدام هذا النص في أي حقل إدخال أدناه.")
                else:
                    st.error("لم يتم التعرف على الصوت. تأكد من وضوح التسجيل.")
    else:
        st.warning("⚠️ streamlit-mic-recorder غير مثبت.")

    enable_web_search = st.checkbox("🌐 تمكين البحث عبر الإنترنت (Tavily)", value=False, key="global_web_search")
    if enable_web_search and not TAVILY_API_KEY:
        st.error("مفتاح Tavily غير موجود. أضف TAVILY_API_KEY في secrets.")

    level = st.selectbox("🏫 الطور التعليمي", list(CURRICULUM.keys()))
    info = CURRICULUM[level]
    grade = st.selectbox("📚 المستوى", info["grades"])
    branch = None
    if info["branches"] and grade in info["branches"]:
        branch = st.selectbox("🎯 الشعبة", list(info["branches"][grade].keys()))
    if info["subjects"]:
        subj_list = info["subjects"].get(grade) or info["subjects"].get("_default", [])
    elif info["branches"] and grade in info["branches"] and branch:
        subj_list = info["branches"][grade][branch]
    else:
        subj_list = []
    subject = (st.selectbox("📖 المادة", subj_list) if subj_list else st.text_input("📖 المادة", key="sb_subject"))

    st.markdown("---")
    st.markdown("**🏫 معلومات المؤسسة**")
    school_name = st.text_input("اسم المتوسطة / الثانوية", placeholder="متوسطة الشهيد...", key="school_name")
    teacher_name = st.text_input("اسم الأستاذ(ة)", placeholder="الأستاذ(ة)...", key="teacher_name")
    wilaya = st.text_input("الولاية", placeholder="الجزائر...", key="wilaya")
    school_year = st.text_input("السنة الدراسية", value="2025/2026", key="syear")

    st.markdown("---")
    st.markdown("**تواصل — DONIA LABS TECH**")
    st.markdown(
        f"""
        <div class="donia-social">
          <a href="{SOCIAL_URL_WHATSAPP}" target="_blank" rel="noopener noreferrer" title="WhatsApp">📱 WA</a>
          <a href="{SOCIAL_URL_LINKEDIN}" target="_blank" rel="noopener noreferrer" title="LinkedIn">💼 in</a>
          <a href="{SOCIAL_URL_FACEBOOK}" target="_blank" rel="noopener noreferrer" title="Facebook">📖 f</a>
          <a href="{SOCIAL_URL_TELEGRAM}" target="_blank" rel="noopener noreferrer" title="Telegram">✈️ TG</a>
        </div>
        """,
        unsafe_allow_html=True,
    )

# ========== HEADER ==========
st.markdown("""
<div class="donia-slogan-bar">
  <span class="donia-slogan-ar">بالعلم نرتقي</span>
  <div class="donia-slogan-divider"></div>
  <span class="donia-slogan-en">Education Uplifts Us</span>
</div>
""", unsafe_allow_html=True)

st.markdown(f"""
<div class="title-card">
    <h1 style="color:#ffffff!important;font-family:'Cairo',sans-serif">🎓 DONIA MIND — المعلم الذكي v5.0</h1>
    <div class="donia-robot-wrap" aria-hidden="true">
      <div class="donia-robot" title="مساعدك التربوي الذكي">
        <svg viewBox="0 0 80 80" xmlns="http://www.w3.org/2000/svg">
          <rect x="15" y="18" width="50" height="44" rx="14" fill="#d5f5e3" stroke="#145a32" stroke-width="2.5"/>
          <line x1="40" y1="18" x2="40" y2="8" stroke="#c0392b" stroke-width="3" stroke-linecap="round"/>
          <circle cx="40" cy="6" r="4" fill="#c0392b">
            <animate attributeName="r" values="4;5.5;4" dur="1.6s" repeatCount="indefinite"/>
            <animate attributeName="opacity" values="1;.55;1" dur="1.6s" repeatCount="indefinite"/>
          </circle>
          <circle cx="31" cy="36" r="6" fill="#145a32"/>
          <circle cx="49" cy="36" r="6" fill="#145a32"/>
          <circle cx="32.5" cy="35" r="2.2" fill="#ffffff">
            <animateTransform attributeName="transform" type="translate" values="0,0;1,0;0,0;-1,0;0,0" dur="3s" repeatCount="indefinite"/>
          </circle>
          <circle cx="50.5" cy="35" r="2.2" fill="#ffffff">
            <animateTransform attributeName="transform" type="translate" values="0,0;1,0;0,0;-1,0;0,0" dur="3s" repeatCount="indefinite"/>
          </circle>
          <path d="M30 52 Q40 60 50 52" stroke="#c0392b" stroke-width="3" fill="none" stroke-linecap="round">
            <animate attributeName="d" values="M30 52 Q40 60 50 52;M30 50 Q40 58 50 50;M30 52 Q40 60 50 52" dur="2.5s" repeatCount="indefinite"/>
          </path>
          <line x1="23" y1="28" x2="23" y2="46" stroke="rgba(20,90,50,.25)" stroke-width="1.2" stroke-dasharray="3 2"/>
          <line x1="57" y1="28" x2="57" y2="46" stroke="rgba(20,90,50,.25)" stroke-width="1.2" stroke-dasharray="3 2"/>
          <ellipse cx="40" cy="68" rx="18" ry="4.5" fill="rgba(39,174,96,.25)"/>
        </svg>
      </div>
    </div>
    <p style="font-family:'Cairo',sans-serif;font-weight:600">
      منصة تعليمية للمنظومة الجزائرية · مذكرات · اختبارات · تنقيط · تحليل · تصحيح · ذكاء مزدوج
    </p>
</div>
""", unsafe_allow_html=True)

st.markdown(f'<div class="welcome-banner">🌟 {WELCOME_MESSAGE_AR}</div>', unsafe_allow_html=True)

render_floating_assistant()

# ========== TABS ==========
(tab_plan, tab_exam, tab_grade, tab_report,
 tab_ex, tab_correct, tab_template, tab_archive, tab_stats) = st.tabs([
    "📝 مذكرة الدرس", "📄 توليد اختبار", "📊 دفتر التنقيط",
    "📈 تحليل النتائج", "✏️ توليد تمرين", "✅ تصحيح أوراق",
    "🧠 تعلم القوالب (RAG)", "🗄️ الأرشيف", "📉 إحصائيات",
])

branch_txt = f" – {branch}" if branch else ""

def _unique_suffix():
    return str(uuid.uuid4()).replace("-", "")[:16]

# ========== TAB 1 — Lesson Plan ==========
with tab_plan:
    st.markdown("### 📝 إعداد المذكرة وفق الصيغة الرسمية الجزائرية")
    st.markdown(
        '<div class="template-box">📋 تُنشأ المذكرة بالهيكل الرسمي: '
        'المعلومات العامة · المورد المعرفي · الكفاءة · '
        'سير الدرس (تهيئة - بناء - استثمار) · التقويم · الواجب المنزلي</div>',
        unsafe_allow_html=True)

    templates = get_all_templates()
    template_options = {f"{tid} - {name}": tid for tid, name, _, _ in templates}
    selected_template_desc = st.selectbox("📚 استخدام قالب محفوظ (اختياري)", ["بدون قالب"] + list(template_options.keys()), key="plan_template_sel")
    template_id = None
    template_text = ""
    if selected_template_desc != "بدون قالب":
        template_id = template_options[selected_template_desc]
        template_text, template_struct = get_template_by_id(template_id)
        st.info(f"تم تحميل القالب: {selected_template_desc} - سيتم استخدام بنيته للتوليد.")

    pm1, pm2 = st.columns(2)
    with pm1:
        plan_lesson = st.text_input("📝 عنوان الدرس / المورد المعرفي:", key="plan_lesson",
                                      placeholder="مثال: القاسم المشترك الأكبر لعددين طبيعيين")
        plan_chapter = st.text_input("📚 الباب / الوحدة:", key="plan_chapter",
                                      placeholder="مثال: الباب الأول – الأعداد الطبيعية")
        plan_domain = st.selectbox("🗂️ الميدان:", ["أنشطة عددية", "أنشطة جبرية", "أنشطة هندسية", "أنشطة إحصائية", "ميدان عام"], key="plan_domain")
        plan_dur = st.selectbox("⏱️ مدة الحصة:", ["50 دقيقة", "1 ساعة", "1.5 ساعة", "2 ساعة"], key="plan_dur")
    with pm2:
        plan_session = st.selectbox("نوع الحصة:", ["درس نظري", "أعمال موجهة", "أعمال تطبيقية", "تقييم تشخيصي", "دعم وعلاج"], key="plan_session")
        plan_prereq = st.text_area("📌 المكتسبات القبلية:", key="plan_prereq", height=70,
                                     placeholder="مثال: القسمة الإقليدية، قواسم عدد طبيعي...")
        plan_tools = st.text_input("🛠️ الوسائل والأدوات:", key="plan_tools",
                                      value="الكتاب المدرسي، المنهاج، الوثيقة المرافقة، دليل الأستاذ، السبورة")
        plan_notes = st.text_area("📌 ملاحظات خاصة:", key="plan_notes", height=70,
                                     placeholder="توجيهات خاصة بالفوج...")
        use_critic = st.checkbox("🔍 تفعيل طبقة النقد البيداغوجي (Dual‑AI)", value=True, key="plan_critic")
        if enable_web_search and TAVILY_API_KEY:
            web_enhance = st.checkbox("🌐 تضمين نتائج بحث من الإنترنت", value=False, key="plan_web")

    if st.button("📝 توليد المذكرة الكاملة بالذكاء الاصطناعي", key="btn_gen_plan"):
        if not GROQ_API_KEY:
            st.warning("⚠️ أضف GROQ_API_KEY في متغيرات البيئة لإكمال التوليد.")
        elif not plan_lesson.strip():
            st.warning("⚠️ أدخل عنوان الدرس / المورد المعرفي لإكمال المذكرة.")
        else:
            web_context = ""
            if enable_web_search and web_enhance:
                with st.spinner("🌐 جلب معلومات إضافية من الإنترنت..."):
                    search_query = f"{subject} {plan_lesson} منهاج جزائري {grade}"
                    web_context = web_search(search_query)
                    if web_context:
                        web_context = f"\nمعلومات إضافية من الإنترنت:\n{web_context}\n"
            template_instruction = ""
            if template_text:
                template_instruction = f"استخدم هذا الهيكل المستخرج من القالب كمرجع أساسي:\n{template_text[:2000]}\n"
            prompt = f"""أنت أستاذ جزائري خبير. أعدّ مذكرة درس رسمية وفق المنهاج الجزائري.

{template_instruction}
المعطيات:
• الطور: {level} | السنة: {grade}{branch_txt}
• المادة: {subject} | الميدان: {plan_domain}
• الباب: {plan_chapter} | الدرس: {plan_lesson}
• نوع الحصة: {plan_session} | المدة: {plan_dur}
• المكتسبات القبلية: {plan_prereq}
{f"• ملاحظات: {plan_notes}" if plan_notes.strip() else ""}
{web_context}

{llm_output_language_clause(subject)}

أعدّ المذكرة بهذا الهيكل:

## الكفاءة الختامية
[اكتب الكفاءة الختامية للوحدة]

## مستوى من الكفاءة
[اكتب مستوى الكفاءة المستهدف في هذه الحصة]

## مرحلة التهيئة (5 دقائق)
[نشاط الاستعداد والتمهيد بأسئلة مراجعة]

## أنشطة بناء الموارد (25-30 دقيقة)
### وضعية تعلمية
[وصف النشاط التعلمي التفصيلي مع أمثلة ومعادلات LaTeX حيث يلزم]

### حوصلة
[الخلاصة والقاعدة التي يصل إليها الأستاذ مع التلاميذ]

## مرحلة إعادة الاستثمار (15 دقيقة)
### حل التمرين
[تمرين تطبيقي مع حله التفصيلي]

## التقويم والإرشادات
[أسئلة تقييمية وتوجيهات للأستاذ أثناء كل مرحلة]

## الواجب المنزلي
[تمارين المنزل مع رقم الصفحة إن أمكن]

## نقد ذاتي
[ملاحظات بيداغوجية لما بعد الحصة]"""
            with st.spinner("📝 جاري إعداد المذكرة..."):
                try:
                    plan_text, critic_report = dual_llm_generate_with_critic(prompt, subject, grade, use_critic=use_critic)
                    if critic_report.get("validated"):
                        st.success(f"✅ تم التحقق من المحتوى (درجة الناقد: {critic_report.get('score',0)}/10)")
                        if critic_report.get("remarks"):
                            st.info(f"📌 ملاحظات الناقد: {critic_report['remarks']}")
                    else:
                        st.warning("⚠️ لم يتم التحقق من المحتوى بواسطة الناقد.")
                    render_with_latex(plan_text)

                    plots = auto_generate_plots(plan_text, subject)
                    for _ptype, _pdata in plots:
                        if _ptype == "plotly":
                            st.plotly_chart(_pdata, use_container_width=True)
                        elif _ptype == "image":
                            st.image(_pdata, use_container_width=True)

                    def extract_section(text, marker):
                        m = re.search(rf'## {marker}[^\n]*\n([\s\S]+?)(?=## |\Z)', text)
                        return m.group(1).strip() if m else ""

                    plan_data = {
                        "school": school_name, "teacher": teacher_name,
                        "grade": f"{grade}{branch_txt}", "domain": plan_domain,
                        "chapter": plan_chapter, "lesson": plan_lesson,
                        "session_type": plan_session, "duration": plan_dur,
                        "subject": subject,
                        "duration_t": "5 د", "duration_b": "25 د", "duration_r": "15 د",
                        "competency": extract_section(plan_text, "مستوى من الكفاءة"),
                        "intro": extract_section(plan_text, "مرحلة التهيئة"),
                        "build": extract_section(plan_text, "أنشطة بناء الموارد"),
                        "reinvest": extract_section(plan_text, "مرحلة إعادة الاستثمار"),
                        "eval": extract_section(plan_text, "التقويم والإرشادات"),
                        "homework": extract_section(plan_text, "الواجب المنزلي"),
                        "self_critique": extract_section(plan_text, "نقد ذاتي"),
                        "prerequisites": plan_prereq, "tools": plan_tools,
                    }

                    db_exec(
                        "INSERT INTO lesson_plans "
                        "(level,grade,subject,lesson,domain,duration,content,created_at) "
                        "VALUES (?,?,?,?,?,?,?,?)",
                        (level, grade, subject, plan_lesson, plan_domain, plan_dur,
                         plan_text, datetime.now().strftime("%Y-%m-%d %H:%M")))
                    st.success("✅ تم حفظ المذكرة")

                    unique_id = _unique_suffix()
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.download_button("📥 نص",
                                           plan_text.encode("utf-8-sig"),
                                           f"مذكرة_{plan_lesson}.txt",
                                           key=f"plan_txt_{unique_id}")
                    with col2:
                        try:
                            pdf_p = generate_lesson_plan_pdf(plan_data)
                            st.download_button("📄 PDF", pdf_p,
                                               f"مذكرة_{plan_lesson}.pdf", "application/pdf",
                                               key=f"plan_pdf_{unique_id}")
                        except Exception as _pe:
                            st.caption(f"⚠️ PDF: {_pe}")
                    with col3:
                        if _DOCX_AVAILABLE:
                            try:
                                docx_p = generate_lesson_plan_docx(plan_data)
                                st.download_button("📝 Word", docx_p,
                                                   f"مذكرة_{plan_lesson}.docx",
                                                   "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                                   key=f"plan_docx_{unique_id}")
                            except Exception as _we:
                                st.caption(f"⚠️ Word: {_we}")
                        else:
                            st.caption("⚠️ python-docx غير مثبت")
                    with col4:
                        xlsx_p = generate_text_excel(plan_text, f"مذكرة: {plan_lesson}",
                                                     {"المادة": subject, "المستوى": grade,
                                                      "الميدان": plan_domain, "الدرس": plan_lesson})
                        st.download_button("📊 Excel", xlsx_p,
                                           f"مذكرة_{plan_lesson}.xlsx",
                                           "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                           key=f"plan_xlsx_{unique_id}")
                    if st.button("💾 حفظ في RAG", key=f"save_rag_plan_{unique_id}"):
                        save_to_rag(plan_text, "lesson_plan", {"subject": subject, "grade": grade, "lesson": plan_lesson})
                        st.success("✅ تم حفظ المذكرة في قاعدة المعرفة RAG")
                except Exception as err:
                    st.error(f"⚠️ تعذر إكمال توليد المذكرة: {err}")

# ========== TAB 2 — Exam Generation ==========
with tab_exam:
    st.markdown("### 📄 توليد ورقة الاختبار وفق النموذج الجزائري الرسمي")
    st.markdown(
        '<div class="template-box">📋 يُنشأ الاختبار بالهيكل الرسمي: '
        'رأس الورقة (المؤسسة، المستوى، المدة) · '
        '4 تمارين بنقاط محددة · وضعية إدماجية 8 نقاط</div>',
        unsafe_allow_html=True)

    templates = get_all_templates()
    template_options = {f"{tid} - {name}": tid for tid, name, _, _ in templates}
    selected_template_desc = st.selectbox("📚 استخدام قالب محفوظ (اختياري)", ["بدون قالب"] + list(template_options.keys()), key="exam_template_sel")
    template_id = None
    template_text = ""
    if selected_template_desc != "بدون قالب":
        template_id = template_options[selected_template_desc]
        template_text, template_struct = get_template_by_id(template_id)
        st.info(f"تم تحميل القالب: {selected_template_desc} - سيتم استخدام بنيته للتوليد.")

    ex1, ex2, ex3 = st.columns(3)
    with ex1:
        exam_semester = st.selectbox("الفصل:", ["الفصل الأول", "الفصل الثاني", "الفصل الثالث"], key="exam_semester")
        exam_duration = st.selectbox("المدة:", ["ساعة واحدة", "ساعتان", "ثلاث ساعات"], key="exam_dur")
    with ex2:
        exam_theme = st.text_input("محاور الاختبار:", key="exam_theme", placeholder="مثال: الجمل, الدوال الخطية, الأعداد الناطقة")
        exam_points = st.text_input("نقاط التمارين:", value="3,3,3,3,8", key="exam_pts", help="مثال: 3,3,3,3,8 (4 تمارين + وضعية إدماجية)")
    with ex3:
        exam_difficulty = st.select_slider("مستوى الصعوبة:", ["سهل", "متوسط", "صعب", "مستوى الشهادة"], key="exam_diff")
        include_integrate = st.checkbox("إضافة وضعية إدماجية", value=True, key="exam_integrate")
        use_critic_exam = st.checkbox("🔍 تفعيل النقد البيداغوجي (Dual‑AI)", value=True, key="exam_critic")
        if enable_web_search and TAVILY_API_KEY:
            exam_web = st.checkbox("🌐 بحث عن نماذج اختبارات من الإنترنت", value=False, key="exam_web")

    exam_notes = st.text_area("ملاحظات وتوجيهات:", key="exam_notes", placeholder="مثلاً: التركيز على الأعداد الناطقة والجذور التربيعية...")

    if st.button("🚀 توليد ورقة الاختبار", key="btn_gen_exam"):
        if not GROQ_API_KEY:
            st.error("⚠️ أضف GROQ_API_KEY")
        else:
            pts = exam_points.split(",")
            pts_desc = " + ".join([f"تمرين {i+1}: {p.strip()} نقاط" for i, p in enumerate(pts[:4])])
            integrate_txt = (f"+ وضعية إدماجية: {pts[4].strip() if len(pts) > 4 else '8'} نقاط" if include_integrate else "")
            web_ctx = ""
            if enable_web_search and exam_web:
                with st.spinner("جلب نماذج من الإنترنت..."):
                    web_ctx = web_search(f"اختبار {subject} {grade} {exam_semester} جزائري")
                    if web_ctx:
                        web_ctx = f"\nمقترحات من الإنترنت:\n{web_ctx}\n"
            template_instruction = ""
            if template_text:
                template_instruction = f"استخدم هذا الهيكل المستخرج من القالب كمرجع أساسي:\n{template_text[:2000]}\n"
            prompt = f"""أنت أستاذ جزائري خبير في إعداد الاختبارات. أعدّ ورقة اختبار رسمية.

{template_instruction}
المعطيات:
• الطور: {level} | المستوى: {grade}{branch_txt}
• المادة: {subject} | {exam_semester}
• المدة: {exam_duration} | الصعوبة: {exam_difficulty}
• المحاور: {exam_theme or subject}
• توزيع النقاط: {pts_desc} {integrate_txt}
• المجموع: 20 نقطة
{f"• ملاحظات: {exam_notes}" if exam_notes.strip() else ""}
{web_ctx}

{llm_output_language_clause(subject)}

أعدّ الاختبار بهذا الهيكل الدقيق:

تمرين 1 :( {pts[0].strip() if pts else '3'} نقاط)
[الأسئلة مرقمة]

تمرين 2 :( {pts[1].strip() if len(pts) > 1 else '3'} نقاط)
[الأسئلة...]

تمرين 3 :( {pts[2].strip() if len(pts) > 2 else '3'} نقاط)
[الأسئلة...]

تمرين 4 :( {pts[3].strip() if len(pts) > 3 else '3'} نقاط)
[الأسئلة...]

{"الوضعية الإدماجية:( " + (pts[4].strip() if len(pts) > 4 else '8') + " نقاط)" if include_integrate else ""}
{"السياق: [سياق واقعي جزائري]" if include_integrate else ""}
{"الجزء الأول: [أسئلة تدريجية...]" if include_integrate else ""}
{"الجزء الثاني: [أسئلة تكملة...]" if include_integrate else ""}
{"انتهى — بالتوفيق والنجاح" if include_integrate else ""}

القواعد الإلزامية:
""" + (
"""• اللغة العربية الفصحى للنصوص التعليمية
• المعادلات بتنسيق LaTeX
• الأسئلة مرقمة ومتدرجة في الصعوبة"""
    if get_pdf_mode_for_subject(subject)[0] else
"""• Use ONLY the target foreign language for all instructional text
• Equations in LaTeX where appropriate
• Numbered questions, progressive difficulty"""
) + """
"""
            with st.spinner("📄 جاري توليد الاختبار..."):
                try:
                    exam_content, critic_report = dual_llm_generate_with_critic(prompt, subject, grade, use_critic=use_critic_exam)
                    if critic_report.get("validated"):
                        st.success(f"✅ تم التحقق (نقاط الناقد: {critic_report.get('score',0)}/10)")
                    st.markdown(
                        f'<div class="feature-card"><h4>📄 {subject} | {grade}{branch_txt} | {exam_semester} | ⏱️ {exam_duration}</h4></div>',
                        unsafe_allow_html=True)
                    render_with_latex(exam_content)
                    plots = auto_generate_plots(exam_content, subject)
                    for _ptype, _pdata in plots:
                        if _ptype == "plotly":
                            st.plotly_chart(_pdata, use_container_width=True)
                        elif _ptype == "image":
                            st.image(_pdata, use_container_width=True)
                    db_exec(
                        "INSERT INTO exams (level,grade,subject,semester,content,created_at) "
                        "VALUES (?,?,?,?,?,?)",
                        (level, grade, subject, exam_semester, exam_content,
                         datetime.now().strftime("%Y-%m-%d %H:%M")))
                    st.success("✅ تم حفظ الاختبار")

                    exam_pdf_data = {
                        "school": school_name, "wilaya": wilaya,
                        "grade": f"{grade}{branch_txt}", "year": school_year,
                        "district": "...", "semester": exam_semester,
                        "subject": subject, "duration": exam_duration,
                        "content": exam_content,
                    }
                    unique_id = _unique_suffix()
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.download_button("📥 نص",
                                           exam_content.encode("utf-8-sig"),
                                           f"اختبار_{subject}_{exam_semester}.txt",
                                           key=f"exam_txt_{unique_id}")
                    with col2:
                        try:
                            pdf_e = generate_exam_pdf(exam_pdf_data)
                            st.download_button("📄 PDF", pdf_e,
                                               f"اختبار_{subject}_{exam_semester}.pdf",
                                               "application/pdf", key=f"exam_pdf_{unique_id}")
                        except Exception as _pe:
                            st.caption(f"⚠️ PDF: {_pe}")
                    with col3:
                        if _DOCX_AVAILABLE:
                            try:
                                docx_e = generate_exam_docx(exam_pdf_data)
                                st.download_button("📝 Word", docx_e,
                                                   f"اختبار_{subject}_{exam_semester}.docx",
                                                   "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                                   key=f"exam_docx_{unique_id}")
                            except Exception as _we:
                                st.caption(f"⚠️ Word: {_we}")
                        else:
                            st.caption("⚠️ python-docx غير مثبت")
                    with col4:
                        xlsx_e = generate_text_excel(exam_content, f"اختبار: {subject}",
                                                     {"المادة": subject, "المستوى": grade,
                                                      "الفصل": exam_semester, "المدة": exam_duration})
                        st.download_button("📊 Excel", xlsx_e,
                                           f"اختبار_{subject}_{exam_semester}.xlsx",
                                           "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                           key=f"exam_xlsx_{unique_id}")
                    if st.button("💾 حفظ في RAG", key=f"save_rag_exam_{unique_id}"):
                        save_to_rag(exam_content, "exam", {"subject": subject, "grade": grade, "semester": exam_semester})
                        st.success("✅ تم حفظ الاختبار في قاعدة المعرفة RAG")
                except Exception as err:
                    st.error(f"❌ {err}")

# ========== TAB 3 — Grade Book ==========
with tab_grade:
    st.markdown("### 📊 دفتر التنقيط الرسمي")
    grade_mode = st.radio("وضع الإدخال:",
                           ["📁 رفع ملف Excel (دفتر موجود)", "✏️ إدخال يدوي"],
                           horizontal=True, key="grade_mode")
    students_data = []
    if grade_mode == "📁 رفع ملف Excel (دفتر موجود)":
        gr_file = st.file_uploader("📁 ارفع ملف دفتر التنقيط:",
                                    type=["xlsx", "xls"], key="gr_upload")
        if gr_file:
            _sheet_names = list_excel_sheet_names(gr_file)
            gr_merge = st.checkbox(
                "دمج جميع أوراق الملف (Sheets) في قائمة واحدة",
                value=False, key="gr_merge_all",
                help="يفيد عند وجود عدة أقسام/أفواج في نفس الملف.")
            gr_sel = None
            if not gr_merge and len(_sheet_names) > 1:
                gr_sel = st.selectbox(
                    "اختر الورقة المراد قراءتها:", _sheet_names, key="gr_sheet_pick")
            elif not gr_merge and len(_sheet_names) == 1:
                gr_sel = _sheet_names[0]
            with st.spinner("جاري قراءة الملف..."):
                try:
                    students_data = parse_grade_book_excel(
                        gr_file, sheet_name=gr_sel, merge_all_sheets=gr_merge)
                    st.success(f"✅ تم قراءة {len(students_data)} تلميذ")
                except Exception as e:
                    st.error(f"خطأ في القراءة: {e}")
    else:
        st.markdown("**أدخل بيانات التلاميذ (اسم، تقويم، فرض، اختبار) — سطر لكل تلميذ:**")
        manual_data = st.text_area("", height=200, key="grade_manual",
            placeholder="أحمد بلعيد, 15, 12, 14\nفاطمة زروق, 18, 17, 19\nعلي حمدي, 10, 8, 11")
        if manual_data.strip():
            for idx, line in enumerate(manual_data.strip().splitlines(), 1):
                parts = [p.strip() for p in line.split(",")]
                if len(parts) >= 4:
                    try:
                        name_parts = parts[0].split()
                        students_data.append({
                            'id': idx,
                            'nom': name_parts[0] if name_parts else parts[0],
                            'prenom': " ".join(name_parts[1:]) if len(name_parts) > 1 else '',
                            'dob': '', 'taqwim': float(parts[1]),
                            'fard': float(parts[2]), 'ikhtibhar': float(parts[3]),
                        })
                    except (ValueError, IndexError):
                        pass
            for s in students_data:
                s['average'] = calc_average(s['taqwim'], s['fard'], s['ikhtibhar'])
                s['apprec'] = get_appreciation(s['average'])

    if students_data:
        gc1, gc2 = st.columns(2)
        with gc1:
            gb_class = st.text_input("اسم القسم:", placeholder="4م1", key="gb_class")
            gb_sem = st.selectbox("الفصل:", ["الفصل الأول", "الفصل الثاني", "الفصل الثالث"], key="gb_sem")
        with gc2:
            gb_subject = st.text_input("المادة:", value=subject, key="gb_subject")
            gb_school = st.text_input("المؤسسة:", value=school_name, key="gb_school")
        df = pd.DataFrame([{
            "الرقم": idx + 1,
            "اللقب": s.get('nom', ''),
            "الاسم": s.get('prenom', ''),
            "الورقة": s.get('sheet_source', ''),
            "تقويم /20": s.get('taqwim', ''),
            "فرض /20": s.get('fard', ''),
            "اختبار /20": s.get('ikhtibhar', ''),
            "المعدل": s.get('average', 0),
            "التقدير": s.get('apprec', '')
        } for idx, s in enumerate(students_data)])
        st.markdown("#### 📋 جدول النتائج")
        st.dataframe(df, use_container_width=True, height=350)
        averages = [s['average'] for s in students_data]
        passed = [a for a in averages if a >= 10]
        a1, a2, a3, a4, a5 = st.columns(5)
        for col, val, lbl, clr in [
            (a1, len(students_data), "عدد التلاميذ", "#667eea"),
            (a2, f"{sum(averages)/max(len(averages),1):.2f}", "معدل القسم", "#764ba2"),
            (a3, f"{max(averages):.2f}" if averages else "—", "أعلى معدل", "#10b981"),
            (a4, f"{min(averages):.2f}" if averages else "—", "أدنى معدل", "#ef4444"),
            (a5, f"{len(passed)}/{len(averages)}", "الناجحون", "#f59e0b"),
        ]:
            with col:
                st.markdown(
                    f'<div class="stat-card"><h2 style="color:{clr}">{val}</h2>'
                    f'<p>{lbl}</p></div>', unsafe_allow_html=True)
        fig = px.bar(df, x="اللقب", y="المعدل", color="التقدير",
            color_discrete_map={
                "ممتاز": "#10b981", "جيد جداً": "#3b82f6", "جيد": "#667eea",
                "مقبول": "#f59e0b", "ضعيف": "#ef4444"},
            title=f"نتائج {gb_class or 'القسم'}", template="plotly_dark")
        fig.add_hline(y=10, line_dash="dash", line_color="red", annotation_text="حد النجاح")
        fig.update_layout(paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)")
        st.plotly_chart(fig, use_container_width=True)
        dg1, dg2 = st.columns(2)
        with dg1:
            if len(set([s.get('sheet_source', gb_class) for s in students_data])) > 1:
                classes_by_sheet = {}
                for s in students_data:
                    sheet = s.get('sheet_source', gb_class)
                    if sheet not in classes_by_sheet:
                        classes_by_sheet[sheet] = []
                    classes_by_sheet[sheet].append(s)
                classes_data = [{"name": sheet, "students": students} for sheet, students in classes_by_sheet.items()]
                xlsx_bytes = generate_multi_sheet_grade_book(
                    classes_data, gb_school or school_name,
                    gb_subject or subject, gb_sem)
                st.download_button(
                    "📊 تحميل دفتر التنقيط (Excel - متعدد الأوراق)", xlsx_bytes,
                    f"دفتر_الأقسام_{gb_sem}.xlsx",
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="dl_grade_xlsx_multi")
            else:
                xlsx_bytes = generate_grade_book_excel(
                    students_data, gb_class or "القسم",
                    gb_subject or subject, gb_sem, gb_school or school_name)
                st.download_button(
                    "📊 تحميل دفتر التنقيط (Excel)", xlsx_bytes,
                    f"دفتر_{gb_class}_{gb_sem}.xlsx",
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="dl_grade_xlsx")
        with dg2:
            if st.button("💾 حفظ في قاعدة البيانات", key="btn_save_grade"):
                db_exec(
                    "INSERT INTO grade_books "
                    "(class_name,subject,semester,data_json,created_at) "
                    "VALUES (?,?,?,?,?)",
                    (gb_class, subject, gb_sem,
                     json.dumps(students_data, ensure_ascii=False),
                     datetime.now().strftime("%Y-%m-%d %H:%M")))
                st.success("✅ تم الحفظ")

# ========== TAB 4 — Report Analysis ==========
with tab_report:
    st.markdown("### 📈 تحليل نتائج الأقسام (تقرير شامل)")
    rep_mode = st.radio("مصدر البيانات:",
        ["📁 رفع ملف Excel", "📋 إدخال يدوي", "📂 من قاعدة البيانات"],
        horizontal=True, key="rep_mode")
    all_classes = []
    if rep_mode == "📁 رفع ملف Excel":
        rep_files = st.file_uploader(
            "📁 ارفع ملفات دفتر التنقيط (يمكن رفع عدة أقسام):",
            type=["xlsx"], accept_multiple_files=True, key="rep_upload")
        rep_merge_sheets = st.checkbox(
            "دمج جميع أوراق كل ملف Excel", value=False, key="rep_merge_all",
            help="عند التفعيل تُقرأ كل الأوراق وتُدمج لكل ملف.")
        rep_sheet_choice = None
        if rep_files and not rep_merge_sheets:
            _sn0 = list_excel_sheet_names(rep_files[0])
            if len(_sn0) > 1:
                rep_sheet_choice = st.selectbox(
                    "الورقة المستخدمة (يُفترض تطابق أسماء الأوراق بين الملفات):",
                    _sn0, key="rep_sheet_pick")
            elif _sn0:
                rep_sheet_choice = _sn0[0]
        if rep_files:
            for f in rep_files:
                try:
                    stus = parse_grade_book_excel(
                        f, sheet_name=rep_sheet_choice, merge_all_sheets=rep_merge_sheets)
                    if stus:
                        cls_name = f.name.replace(".xlsx", "").replace("_", " ")
                        all_classes.append(build_class_stats(stus, cls_name))
                except Exception as e:
                    st.warning(f"خطأ في {f.name}: {e}")
    elif rep_mode == "📋 إدخال يدوي":
        st.caption("أدخل بيانات كل قسم (اسم القسم, عدد الناجحين, المعدل, المجموع):")
        rep_text = st.text_area("", height=150, key="rep_manual",
            placeholder="4م1, 13, 8.07, 42\n4م2, 14, 8.86, 41\n4م3, 18, 10.5, 40")
        for line in (rep_text or "").strip().splitlines():
            parts = [p.strip() for p in line.split(",")]
            if len(parts) >= 4:
                try:
                    total = int(parts[3])
                    passed_n = int(parts[1])
                    avg = float(parts[2])
                    all_classes.append({
                        "name": parts[0], "total": total,
                        "avg": avg, "max": 20.0, "min": 0.0,
                        "pass_rate": passed_n / max(total, 1) * 100,
                        "distribution": {}, "top5": [], "students": [],
                    })
                except (ValueError, ZeroDivisionError):
                    pass
    else:
        saved = db_exec(
            "SELECT * FROM grade_books ORDER BY created_at DESC LIMIT 20",
            fetch=True) or []
        if not saved:
            st.info("لا توجد بيانات محفوظة بعد.")
        else:
            for row in saved:
                try:
                    rid, cname, sub, sem, data_j, created = row
                    stus = json.loads(data_j)
                    if stus:
                        all_classes.append(build_class_stats(stus, cname))
                except Exception:
                    pass
    if all_classes:
        rep_subject = st.text_input("المادة:", value=subject, key="rep_subj")
        rep_semester = st.selectbox("الفصل:", ["الفصل الأول", "الفصل الثاني", "الفصل الثالث"], key="rep_sem")
        df_cls = pd.DataFrame([{
            "القسم": c['name'], "المعدل": round(c['avg'], 2),
            "نسبة النجاح": round(c['pass_rate'], 1), "عدد التلاميذ": c['total']
        } for c in all_classes])
        ch1, ch2 = st.columns(2)
        with ch1:
            fig1 = px.bar(df_cls, x="القسم", y="المعدل", color="القسم",
                          title="مقارنة معدلات الأقسام", template="plotly_dark")
            fig1.update_layout(paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)")
            st.plotly_chart(fig1, use_container_width=True)
        with ch2:
            fig2 = px.bar(df_cls, x="القسم", y="نسبة النجاح", color="القسم",
                          title="مقارنة نسب النجاح %", template="plotly_dark")
            fig2.add_hline(y=50, line_dash="dash", line_color="red")
            fig2.update_layout(paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)")
            st.plotly_chart(fig2, use_container_width=True)
        st.dataframe(df_cls, use_container_width=True)
        for cls in all_classes:
            with st.expander(f"📊 تفاصيل القسم {cls['name']}"):
                st.markdown(
                    f'<div class="template-box">'
                    f'عدد التلاميذ: <b>{cls["total"]}</b> &nbsp;|&nbsp;'
                    f'المعدل: <b>{safe_f(cls["avg"])}</b> &nbsp;|&nbsp;'
                    f'أعلى: <b>{safe_f(cls["max"])}</b> &nbsp;|&nbsp;'
                    f'أدنى: <b>{safe_f(cls["min"])}</b> &nbsp;|&nbsp;'
                    f'نسبة النجاح: <b>{safe_f(cls["pass_rate"], ".1f")}%</b>'
                    f'</div>', unsafe_allow_html=True)
                if cls.get('top5'):
                    top_df = pd.DataFrame(cls['top5'])
                    top_df.index = range(1, len(top_df) + 1)
                    st.caption("أفضل 5 تلاميذ:")
                    st.dataframe(top_df, use_container_width=True)
                if cls.get('distribution'):
                    dist_df = pd.DataFrame([cls['distribution']])
                    st.caption("توزيع الدرجات:")
                    st.dataframe(dist_df, use_container_width=True)
        if GROQ_API_KEY and st.button("🤖 توليد التقرير البيداغوجي بالذكاء الاصطناعي", key="btn_rep_ai"):
            summary = "\n".join([
                f"القسم {c['name']}: معدل={safe_f(c['avg'])}, نجاح={safe_f(c['pass_rate'],'.1f')}%, عدد={c['total']}"
                for c in all_classes])
            prompt_rep = f"""أنت مستشار بيداغوجي جزائري خبير. حلّل النتائج التالية:
{summary}
المادة: {rep_subject} | {rep_semester} | المستوى: {grade}{branch_txt}

{llm_output_language_clause(rep_subject)}

قدّم تقريراً شاملاً يتضمن:
1. التشخيص العام للمستوى
2. مقارنة بين الأقسام (نقاط القوة والضعف)
3. الفئات التي تحتاج دعماً
4. توصيات بيداغوجية محددة لكل قسم
5. خطة علاجية مقترحة
6. مقترحات للأستاذ لتطوير أدائه"""
            with st.spinner("🧠 جاري التحليل البيداغوجي..."):
                try:
                    ai_analysis, _ = dual_llm_generate_with_critic(prompt_rep, rep_subject, grade, use_critic=False)
                    st.markdown("---")
                    st.markdown("#### 🤖 التقرير البيداغوجي")
                    render_with_latex(ai_analysis)
                    report_data = {
                        "school": school_name, "subject": rep_subject,
                        "semester": rep_semester, "classes": all_classes,
                        "ai_analysis": ai_analysis,
                    }
                    st.session_state.stored_report_data = report_data
                    pdf_rep = generate_report_pdf(report_data)
                    unique_id = _unique_suffix()
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.download_button("📥 تحميل نص التقرير",
                                           ai_analysis.encode("utf-8-sig"),
                                           f"تقرير_نتائج_{rep_semester}.txt",
                                           key=f"rep_txt_{unique_id}")
                    with col2:
                        st.download_button("📄 تحميل التقرير الكامل PDF", pdf_rep,
                                           f"تقرير_نتائج_{rep_semester}.pdf",
                                           "application/pdf", key=f"rep_pdf_{unique_id}")
                    with col3:
                        if _DOCX_AVAILABLE:
                            docx_rep = generate_report_docx(report_data)
                            st.download_button("📝 تحميل Word (.docx)", docx_rep,
                                               f"تقرير_نتائج_{rep_semester}.docx",
                                               "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                               key=f"rep_docx_{unique_id}")
                        else:
                            st.caption("⚠️ python-docx غير مثبت")
                    with col4:
                        if st.button("💾 حفظ في RAG", key=f"save_rag_report_{unique_id}"):
                            save_to_rag(ai_analysis, "report", {"subject": rep_subject, "semester": rep_semester})
                            st.success("✅ تم حفظ التقرير في قاعدة المعرفة RAG")
                except Exception as e:
                    st.error(str(e))
        else:
            report_data = {
                "school": school_name, "subject": rep_subject,
                "semester": rep_semester, "classes": all_classes, "ai_analysis": "",
            }
            pdf_rep = generate_report_pdf(report_data)
            unique_id = _unique_suffix()
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.download_button("📥 تحميل نص التقرير",
                                   "لا يوجد تحليل ذكي".encode("utf-8-sig"),
                                   f"تقرير_نتائج_{rep_semester}.txt",
                                   key=f"rep_txt2_{unique_id}")
            with col2:
                st.download_button("📄 تحميل التقرير PDF", pdf_rep,
                                   "تقرير_نتائج.pdf", "application/pdf",
                                   key=f"rep_pdf2_{unique_id}")
            with col3:
                if _DOCX_AVAILABLE:
                    docx_rep2 = generate_report_docx(report_data)
                    st.download_button("📝 تحميل Word (.docx)", docx_rep2,
                                       "تقرير_نتائج.docx",
                                       "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                       key=f"rep_docx2_{unique_id}")
                else:
                    st.caption("⚠️ python-docx غير مثبت")
            with col4:
                if st.button("💾 حفظ في RAG", key=f"save_rag_report_static_{unique_id}"):
                    save_to_rag(json.dumps(all_classes), "report_static", {"subject": rep_subject, "semester": rep_semester})
                    st.success("✅ تم حفظ البيانات في قاعدة المعرفة RAG")

# ========== TAB 5 — Exercise Generation ==========
with tab_ex:
    st.markdown("### ✏️ توليد تمرين مع الحل التفصيلي")
    c1, c2, c3 = st.columns([4, 1, 1])
    with c1:
        lesson = st.text_input("📝 عنوان الدرس:", key="ex_lesson", placeholder="مثال: الانقسام المنصف، المعادلات التفاضلية…")
    with c2:
        num_ex = st.number_input("عدد التمارين", 1, 5, 1, key="ex_num")
    with c3:
        ex_type = st.selectbox("النوع", ["تمرين تطبيقي", "مسألة", "سؤال إشكالي", "فرض محروس"], key="ex_type")
    difficulty = st.select_slider("⚡ مستوى الصعوبة", ["سهل جداً", "سهل", "متوسط", "صعب", "مستوى بكالوريا"], key="ex_difficulty")
    extra = st.text_area("📌 تعليمات إضافية:", placeholder="أي توجيهات خاصة…", key="ex_extra")
    if st.button("🚀 توليد التمرين والحل التفصيلي", key="btn_gen_ex"):
        if not GROQ_API_KEY:
            st.error("⚠️ أضف GROQ_API_KEY")
        elif not lesson.strip():
            st.warning("⚠️ أدخل عنوان الدرس")
        else:
            prompt = f"""أنت أستاذ جزائري خبير. صمم {num_ex} {ex_type}.

• الطور: {level} | السنة: {grade}{branch_txt}
• المادة: {subject} | الدرس: {lesson} | الصعوبة: {difficulty}
{f"• ملاحظات: {extra}" if extra.strip() else ""}

{llm_output_language_clause(subject)}

الهيكل المطلوب:
## التمرين
[المعطيات والمطلوب بوضوح]

## الحل المفصل
[خطوات مرقمة]

## ملاحظات للأستاذ
[توجيهات بيداغوجية]"""
            with st.spinner("🧠 جاري التوليد…"):
                try:
                    llm = get_llm(DEFAULT_GROQ_MODEL, GROQ_API_KEY)
                    res_text = call_llm(llm, prompt)
                    render_with_latex(res_text)
                    plots = auto_generate_plots(res_text, subject)
                    for _ptype, _pdata in plots:
                        if _ptype == "plotly":
                            st.plotly_chart(_pdata, use_container_width=True)
                        elif _ptype == "image":
                            st.image(_pdata, use_container_width=True)
                    db_exec(
                        "INSERT INTO exercises "
                        "(level,grade,branch,subject,lesson,ex_type,difficulty,content,created_at) "
                        "VALUES (?,?,?,?,?,?,?,?,?)",
                        (level, grade, branch or "", subject, lesson, ex_type,
                         difficulty, res_text, datetime.now().strftime("%Y-%m-%d %H:%M")))
                    st.success("✅ تم الحفظ")
                    unique_id = _unique_suffix()
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.download_button("📥 نص", res_text.encode("utf-8-sig"),
                                           f"{lesson}.txt", key=f"ex_txt_{unique_id}")
                    with col2:
                        try:
                            rtl, _ = get_pdf_mode_for_subject(subject)
                            pdf_ex = generate_simple_pdf(res_text, lesson, f"{subject} | {grade}", rtl=rtl)
                            st.download_button("📄 PDF", pdf_ex, f"{lesson}.pdf",
                                               "application/pdf", key=f"ex_pdf_{unique_id}")
                        except Exception as _pe:
                            st.caption(f"⚠️ PDF: {_pe}")
                    with col3:
                        if _DOCX_AVAILABLE:
                            try:
                                ex_docx_data = {
                                    "school": school_name, "teacher": teacher_name,
                                    "subject": subject, "grade": f"{grade}{branch_txt}",
                                    "lesson": lesson, "domain": "تمارين",
                                    "duration": "غير محدد", "content": res_text
                                }
                                docx_ex = generate_lesson_plan_docx(ex_docx_data)
                                st.download_button("📝 Word", docx_ex,
                                                   f"{lesson}.docx",
                                                   "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                                   key=f"ex_docx_{unique_id}")
                            except Exception as _we:
                                st.caption(f"⚠️ Word: {_we}")
                        else:
                            st.caption("⚠️ python-docx غير مثبت")
                    with col4:
                        xlsx_ex = generate_text_excel(res_text, f"تمرين: {lesson}",
                                                      {"المادة": subject, "المستوى": grade,
                                                       "النوع": ex_type, "الصعوبة": difficulty})
                        st.download_button("📊 Excel", xlsx_ex,
                                           f"{lesson}.xlsx",
                                           "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                           key=f"ex_xlsx_{unique_id}")
                    if st.button("💾 حفظ في RAG", key=f"save_rag_ex_{unique_id}"):
                        save_to_rag(res_text, "exercise", {"subject": subject, "grade": grade, "lesson": lesson})
                        st.success("✅ تم حفظ التمرين في قاعدة المعرفة RAG")
                except Exception as err:
                    st.error(f"❌ {err}")


# ══════════════════════════════════════════════════
# GEOMETRY VISUALIZER — appended to Exercise Tab
# ══════════════════════════════════════════════════
    st.markdown("---")
    st.markdown("### 📐 مُوَلِّد الأشكال الهندسية التفاعلي")
    st.markdown(
        '<div class="template-box">🔷 ارسم أي شكل هندسي وأحسب مساحته ومحيطه تلقائياً</div>',
        unsafe_allow_html=True)
    g1, g2 = st.columns([1, 2])
    with g1:
        geo_shape = st.selectbox("الشكل:", [
            "دائرة", "مستطيل", "مثلث", "مثلث قائم",
            "متوازي أضلاع", "شبه منحرف"], key="geo_shape")
        _gp = {}
        if geo_shape == "دائرة":
            _gp['r'] = st.number_input("r", 0.5, 50.0, 3.0, key="gr")
            _ga = 3.14159 * _gp['r']**2; _gpe = 2 * 3.14159 * _gp['r']
        elif geo_shape == "مستطيل":
            _gp['w'] = st.number_input("L", 0.5, 50.0, 6.0, key="gw")
            _gp['h'] = st.number_input("l", 0.5, 50.0, 4.0, key="gh")
            _ga = _gp['w']*_gp['h']; _gpe = 2*(_gp['w']+_gp['h'])
        elif geo_shape == "مثلث":
            _gp['a'] = st.number_input("القاعدة a", 0.5, 50.0, 5.0, key="gta")
            _gp['h'] = st.number_input("الارتفاع h", 0.5, 50.0, 4.0, key="gth")
            _ga = 0.5*_gp['a']*_gp['h']; _gpe = 0
        elif geo_shape == "مثلث قائم":
            _gp['a'] = st.number_input("a", 0.5, 50.0, 3.0, key="grta")
            _gp['b'] = st.number_input("b", 0.5, 50.0, 4.0, key="grtb")
            _hyp = np.sqrt(_gp['a']**2 + _gp['b']**2)
            _ga = 0.5*_gp['a']*_gp['b']; _gpe = _gp['a']+_gp['b']+_hyp
            st.info(f"الوتر = {_hyp:.4f}")
        elif geo_shape == "متوازي أضلاع":
            _gp['b'] = st.number_input("b", 0.5, 50.0, 6.0, key="gpb")
            _gp['h'] = st.number_input("h", 0.5, 50.0, 3.0, key="gph")
            _gp['skew'] = st.number_input("ميل", 0.1, 10.0, 1.5, key="gsk")
            _ga = _gp['b']*_gp['h']; _gpe = 2*(_gp['b'] + np.sqrt(_gp['h']**2+_gp['skew']**2))
        elif geo_shape == "شبه منحرف":
            _gp['a'] = st.number_input("a", 0.5, 50.0, 6.0, key="gtra")
            _gp['b'] = st.number_input("b", 0.5, 50.0, 3.0, key="gtrb")
            _gp['h'] = st.number_input("h", 0.5, 50.0, 3.0, key="gtrh")
            _ga = 0.5*(_gp['a']+_gp['b'])*_gp['h']
            _gpe = _gp['a']+_gp['b']+2*np.sqrt(_gp['h']**2+((_gp['a']-_gp['b'])/2)**2)
        else:
            _ga = _gpe = 0.0
        _shape_code = {"دائرة":"circle","مستطيل":"rectangle","مثلث":"triangle",
                       "مثلث قائم":"right_triangle","متوازي أضلاع":"parallelogram",
                       "شبه منحرف":"trapezoid"}
        st.markdown(
            f'<div class="success-box">📐 المساحة = <b>{_ga:.4f}</b> وحدة²<br>'
            f'📏 المحيط ≈ <b>{_gpe:.4f}</b> وحدة</div>',
            unsafe_allow_html=True)
        if st.button("🎨 رسم الشكل", key="btn_geo"):
            try:
                _img = generate_geometry_figure(_shape_code[geo_shape], _gp)
                st.session_state["_geo_img"]  = _img
                st.session_state["_geo_name"] = geo_shape
            except Exception as _ge:
                st.error(f"خطأ: {_ge}")
    with g2:
        if st.session_state.get("_geo_img"):
            st.image(st.session_state["_geo_img"],
                     caption=st.session_state.get("_geo_name",""),
                     use_container_width=True)
            st.download_button("⬇️ PNG", st.session_state["_geo_img"],
                               f"شكل_{st.session_state.get('_geo_name','geo')}.png",
                               "image/png", key="dl_geo")
        else:
            st.info("اضغط 'رسم الشكل' لعرضه.")

    st.markdown("---")
    st.markdown("### 📈 رسم الدوال الرياضية")
    fp1, fp2 = st.columns([1, 2])
    with fp1:
        _fx = st.text_input("f(x) =", value="x**2 - 4", key="fx_expr",
                             help="مثال: sin(x), 2*x+3, x**3-x, sqrt(abs(x))")
        _xmin = st.number_input("x الأدنى", value=-10.0, key="fx_xmin")
        _xmax = st.number_input("x الأعلى", value=10.0,  key="fx_xmax")
        if st.button("📊 رسم الدالة", key="btn_fx"):
            try:
                _fimg = generate_function_plot(_fx, (_xmin, _xmax), f"f(x)={_fx}")
                if _fimg:
                    st.session_state["_fx_img"]   = _fimg
                    st.session_state["_fx_label"] = _fx
                else:
                    st.error("تعذر رسم الدالة — راجع الصياغة.")
            except Exception as _fe:
                st.error(f"خطأ: {_fe}")
    with fp2:
        if st.session_state.get("_fx_img"):
            st.image(st.session_state["_fx_img"],
                     caption=f"f(x) = {st.session_state.get('_fx_label','')}",
                     use_container_width=True)
            st.download_button("⬇️ PNG", st.session_state["_fx_img"],
                               "دالة_رياضية.png", "image/png", key="dl_fx")
        else:
            st.info("أدخل دالة واضغط 'رسم الدالة'.")


# ========== TAB 6 — Correction ==========
with tab_correct:
    st.markdown("### ✅ تصحيح أوراق الاختبار")
    correct_mode = st.radio("وضع التصحيح:",
                             ["📝 إدخال نصي", "📋 التحقق من إجابة وفق نموذج الحل",
                              "📷 صورة ورقة (كاميرا أو ملف)"],
                             horizontal=True, key="correct_mode")
    cc1, cc2 = st.columns(2)
    with cc1:
        student_name = st.text_input("اسم الطالب:", key="corr_name", placeholder="اختياري")
        exam_subj = st.text_input("المادة:", value=subject, key="corr_subject")
    with cc2:
        total_marks = st.number_input("العلامة الكاملة:", 10, 100, 20, key="corr_total")
        correct_style = st.selectbox("أسلوب التصحيح:",
                                      ["تصحيح شامل مع تعليقات", "تصحيح مختصر",
                                       "تحديد الأخطاء فقط"], key="corr_style")
    model_answer = st.text_area("✍️ الحل النموذجي / السؤال:", height=120,
                                   key="corr_model_ans",
                                   placeholder="أدخل السؤال أو الحل النموذجي…")
    if correct_mode == "📷 صورة ورقة (كاميرا أو ملف)":
        st.markdown("**معاينة الصورة قبل المعالجة**")
        img_col1, img_col2 = st.columns(2)
        with img_col1:
            try:
                cam_shot = st.camera_input("📷 الكاميرا المباشرة", key="corr_camera")
            except Exception as cam_err:
                st.error(f"⚠️ تعذر الوصول إلى الكاميرا: {cam_err}. تأكد من منح التطبيق صلاحية الوصول إلى الكاميرا (HTTPS مطلوب).")
                cam_shot = None
        with img_col2:
            up_img = st.file_uploader("📁 رفع صورة (PNG / JPG / JPEG / WEBP)",
                                       type=["png", "jpg", "jpeg", "webp"],
                                       key="corr_file_img")
        preview_bytes = None
        if cam_shot is not None:
            preview_bytes = cam_shot.getvalue()
            st.image(cam_shot, caption="معاينة — الكاميرا", use_container_width=True)
        elif up_img is not None:
            preview_bytes = up_img.read()
            st.image(preview_bytes, caption="معاينة — الملف", use_container_width=True)
        if preview_bytes and st.button("🔍 استخراج النص من الصورة (OCR)", key="btn_ocr"):
            ocr_extra = ocr_answer_sheet_image(preview_bytes)
            if ocr_extra.strip():
                st.session_state["corr_student_ans"] = ocr_extra
                st.success("✅ تم استخراج نص من الصورة — يمكنك تعديله في الحقل أدناه.")
                st.rerun()
            else:
                st.warning("⚠️ لم يُستخرج نص (ثبّت pytesseract و Tesseract، أو انسخ النص يدوياً).")
    ta_h = 160 if correct_mode == "📷 صورة ورقة (كاميرا أو ملف)" else 120
    ph = (
        "الصق إجابة الطالب أو استخدم الاستخراج من الصورة…"
        if correct_mode == "📷 صورة ورقة (كاميرا أو ملف)"
        else "انسخ إجابة الطالب هنا…"
    )
    student_answer = st.text_area(
        "📄 إجابة الطالب:", height=ta_h, key="corr_student_ans", placeholder=ph)
    if st.button("✅ تصحيح الإجابة", key="btn_correct"):
        if not GROQ_API_KEY:
            st.error("⚠️ أضف GROQ_API_KEY")
        elif not student_answer.strip():
            st.warning("⚠️ أدخل إجابة الطالب")
        else:
            prompt_corr = f"""أنت أستاذ جزائري خبير. صحّح إجابة الطالب بأسلوب: {correct_style}

المادة: {exam_subj} | العلامة الكاملة: {total_marks}/20
الحل النموذجي: {model_answer or 'غير محدد — قيّم من حيث المنطق العلمي'}
إجابة الطالب: {student_answer}

## التقييم الكلي
العلامة المقترحة: X/{total_marks}
المستوى: [ممتاز/جيد جداً/جيد/مقبول/ضعيف]

## نقاط القوة

## الأخطاء والنواقص

## التوصيات للطالب

## ملاحظة للأستاذ"""
            with st.spinner("🔍 جاري التصحيح…"):
                try:
                    llm = get_llm(DEFAULT_GROQ_MODEL, GROQ_API_KEY)
                    correction = call_llm(llm, prompt_corr)
                    render_with_latex(correction)
                    m = re.search(r'(\d+(?:\.\d+)?)\s*/' + str(total_marks), correction)
                    gv = float(m.group(1)) if m else 0.0
                    db_exec(
                        "INSERT INTO corrections "
                        "(student_name,subject,grade_value,total,feedback,created_at) "
                        "VALUES (?,?,?,?,?,?)",
                        (student_name or "مجهول", exam_subj, gv, total_marks,
                         correction, datetime.now().strftime("%Y-%m-%d %H:%M")))
                    st.success(f"✅ العلامة: {gv}/{total_marks}")
                    unique_id = _unique_suffix()
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.download_button("📥 نص",
                                           correction.encode("utf-8-sig"),
                                           f"تصحيح_{student_name or 'طالب'}.txt",
                                           key=f"corr_txt_{unique_id}")
                    with col2:
                        try:
                            rtl, _ = get_pdf_mode_for_subject(exam_subj)
                            pdf_c = generate_simple_pdf(
                                correction, f"تصحيح: {student_name or 'طالب'}", exam_subj, rtl=rtl)
                            st.download_button("📄 PDF", pdf_c,
                                               f"تصحيح_{student_name or 'طالب'}.pdf",
                                               "application/pdf", key=f"corr_pdf_{unique_id}")
                        except Exception as _pe:
                            st.caption(f"⚠️ PDF: {_pe}")
                    with col3:
                        if _DOCX_AVAILABLE:
                            try:
                                corr_docx_data = {
                                    "school": school_name, "teacher": teacher_name,
                                    "subject": exam_subj, "grade": grade,
                                    "lesson": f"تصحيح {student_name or 'طالب'}",
                                    "domain": "تصحيح", "duration": "غير محدد",
                                    "content": correction
                                }
                                docx_corr = generate_lesson_plan_docx(corr_docx_data)
                                st.download_button("📝 Word", docx_corr,
                                                   f"تصحيح_{student_name or 'طالب'}.docx",
                                                   "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                                   key=f"corr_docx_{unique_id}")
                            except Exception as _we:
                                st.caption(f"⚠️ Word: {_we}")
                        else:
                            st.caption("⚠️ python-docx غير مثبت")
                    with col4:
                        xlsx_c = generate_text_excel(correction,
                                                     f"تصحيح: {student_name or 'طالب'}",
                                                     {"الطالب": student_name or "مجهول",
                                                      "المادة": exam_subj,
                                                      "العلامة": f"{gv}/{total_marks}"})
                        st.download_button("📊 Excel", xlsx_c,
                                           f"تصحيح_{student_name or 'طالب'}.xlsx",
                                           "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                           key=f"corr_xlsx_{unique_id}")
                    if st.button("💾 حفظ في RAG", key=f"save_rag_corr_{unique_id}"):
                        save_to_rag(correction, "correction", {"student": student_name or "مجهول", "subject": exam_subj})
                        st.success("✅ تم حفظ التصحيح في قاعدة المعرفة RAG")
                except Exception as err:
                    st.error(f"❌ {err}")

# ========== TAB 7 — Neural Template Learning ==========
with tab_template:
    st.markdown("### 🧠 تعلم القوالب (RAG System)")
    st.markdown(
        '<div class="template-box">📚 ارفع ملف PDF أو صورة (قالب اختبار، مذكرة، دفتر تنقيط) '
        'ليقوم الذكاء الاصطناعي باستخراج هيكله وتخزينه كقالب يمكن استخدامه لاحقاً في التوليد.</div>',
        unsafe_allow_html=True)
    
    template_file = st.file_uploader("📂 رفع قالب (PDF أو صورة)", type=["pdf", "png", "jpg", "jpeg", "webp"], key="template_upload")
    template_name = st.text_input("اسم القالب (مثال: قالب اختبار رياضيات)", key="template_name", placeholder="اسم مميز للقالب")
    template_type = st.selectbox("نوع القالب", ["اختبار", "مذكرة درس", "تمرين", "دفتر تنقيط"], key="template_type")
    
    if template_file and template_name:
        file_bytes = template_file.read()
        raw_text = ""
        if template_file.type == "application/pdf":
            with st.spinner("جاري استخراج النص من PDF..."):
                raw_text = extract_text_from_pdf(file_bytes)
        else:
            with st.spinner("جاري استخراج النص من الصورة (OCR)..."):
                raw_text = extract_text_from_image(file_bytes)
        
        if raw_text.strip():
            st.success(f"تم استخراج {len(raw_text)} حرف من القالب.")
            with st.expander("معاينة النص المستخرج"):
                st.text(raw_text[:1000] + ("..." if len(raw_text) > 1000 else ""))
            
            if st.button("تحليل هيكل القالب بالذكاء الاصطناعي وحفظه", key="btn_analyze_template"):
                with st.spinner("AI يقوم بتحليل الهيكل..."):
                    try:
                        structure = analyze_template_structure(raw_text, template_type)
                        if "error" not in structure:
                            save_template(template_name, template_type, raw_text, structure)
                            st.success(f"✅ تم حفظ القالب '{template_name}' بنجاح!")
                            st.json(structure)
                        else:
                            st.error(f"فشل التحليل: {structure.get('error')}")
                    except Exception as _tmpl_err:
                        st.error(f"❌ خطأ في تحليل القالب: {_tmpl_err}")
                        try:
                            fallback_struct = {"type": "unknown", "sections": [], "metadata": {},
                                               "key_phrases": [], "suggested_prompt_template": raw_text[:500]}
                            save_template(template_name, template_type, raw_text, fallback_struct)
                            st.warning("⚠️ تم الحفظ بهيكل افتراضي بسبب فشل التحليل.")
                        except Exception as _save_err:
                            st.error(f"❌ فشل الحفظ أيضاً: {_save_err}")
        else:
            st.warning("⚠️ لم يتم استخراج نص من الملف. تأكد من أن الملف يحتوي على نص واضح.")
    
    st.markdown("---")
    st.markdown("### قوالب محفوظة")
    templates = get_all_templates()
    if templates:
        for tid, name, ttype, created in templates:
            with st.expander(f"📌 {name} - {ttype} (تم الحفظ: {created})"):
                if st.button("حذف القالب", key=f"del_template_{tid}"):
                    db_exec("DELETE FROM templates WHERE id=?", (tid,))
                    st.rerun()
    else:
        st.info("لا توجد قوالب محفوظة بعد. ارفع قالباً أعلاه.")

# ========== TAB 8 — Archive ==========
with tab_archive:
    st.markdown("### 🗄️ الأرشيف الشامل")
    arch_tabs = st.tabs(["📚 التمارين", "📝 المذكرات", "📄 الاختبارات", "✅ التصحيحات"])
    with arch_tabs[0]:
        search_q = st.text_input("🔍 بحث:", key="db_search", placeholder="ابحث بعنوان أو مادة…")
        exercises = db_exec(
            "SELECT * FROM exercises WHERE lesson LIKE ? OR subject LIKE ? "
            "ORDER BY created_at DESC",
            (f"%{search_q}%", f"%{search_q}%"), fetch=True) or []
        st.caption(f"النتائج: {len(exercises)}")
        for idx, ex in enumerate(exercises, 1):
            ex_id, lv, gr, br, sub, les, xt, diff, cont, created = ex
            with st.expander(f"📚 {les} | {sub} | {gr} | {diff} | 🕒 {created}"):
                st.markdown(f'<div class="result-box">{cont[:400]}…</div>', unsafe_allow_html=True)
                unique_id = _unique_suffix()
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.download_button("📥 نص", cont.encode("utf-8-sig"),
                                       f"{les}.txt", key=f"arch_txt_{ex_id}_{unique_id}")
                with col2:
                    rtl, _ = get_pdf_mode_for_subject(sub)
                    px2 = generate_simple_pdf(cont, les, rtl=rtl)
                    st.download_button("📄 PDF", px2, f"{les}.pdf",
                                       "application/pdf", key=f"arch_pdf_{ex_id}_{unique_id}")
                with col3:
                    if _DOCX_AVAILABLE:
                        ex_docx_data = {
                            "school": school_name, "teacher": teacher_name,
                            "subject": sub, "grade": gr,
                            "lesson": les, "domain": "تمارين",
                            "duration": "غير محدد", "content": cont
                        }
                        docx_ex = generate_lesson_plan_docx(ex_docx_data)
                        st.download_button("📝 Word", docx_ex, f"{les}.docx",
                                           "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                           key=f"arch_docx_{ex_id}_{unique_id}")
                    else:
                        st.caption("⚠️ Word غير متاح")
                if st.button("🗑️ حذف", key=f"del_{ex_id}"):
                    db_exec("DELETE FROM exercises WHERE id=?", (ex_id,))
                    st.rerun()
    with arch_tabs[1]:
        plans = db_exec("SELECT * FROM lesson_plans ORDER BY created_at DESC", fetch=True) or []
        for p in plans:
            try:
                if p is None or not isinstance(p, (tuple, list)):
                    continue
                if len(p) < 8:
                    continue
                row = list(p) + [None] * max(0, 9 - len(p))
                pid, lv, gr, sub, les, dom, dur, cont, created = row[:9]
                les = "بدون عنوان" if les is None else str(les)
                sub = "" if sub is None else str(sub)
                gr = "" if gr is None else str(gr)
                dom = "" if dom is None else str(dom)
                cont = "" if cont is None else str(cont)
                created = "" if created is None else str(created)
            except Exception:
                continue
            with st.expander(f"📝 {les} | {sub} | {gr} | {dom} | 🕒 {created}"):
                st.markdown(f'<div class="result-box">{cont[:350]}…</div>', unsafe_allow_html=True)
                unique_id = _unique_suffix()
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.download_button("📥 نص", cont.encode("utf-8-sig"),
                                       f"مذكرة_{les}.txt", key=f"plan_txt_{pid}_{unique_id}")
                with col2:
                    rtl, _ = get_pdf_mode_for_subject(sub)
                    ppdf = generate_simple_pdf(cont, f"مذكرة: {les}", f"{sub} | {gr}", rtl=rtl)
                    st.download_button("📄 PDF", ppdf, f"مذكرة_{les}.pdf",
                                       "application/pdf", key=f"plan_pdf_{pid}_{unique_id}")
                with col3:
                    if _DOCX_AVAILABLE:
                        plan_docx_data = {
                            "school": school_name, "teacher": teacher_name,
                            "subject": sub, "grade": gr,
                            "lesson": les, "domain": dom,
                            "duration": dur, "content": cont
                        }
                        docx_plan = generate_lesson_plan_docx(plan_docx_data)
                        st.download_button("📝 Word", docx_plan, f"مذكرة_{les}.docx",
                                           "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                           key=f"plan_docx_{pid}_{unique_id}")
                    else:
                        st.caption("⚠️ Word غير متاح")
    with arch_tabs[2]:
        exams = db_exec("SELECT * FROM exams ORDER BY created_at DESC", fetch=True) or []
        for ex in exams:
            eid, lv, gr, sub, sem, cont, created = ex
            with st.expander(f"📄 {sub} | {gr} | {sem} | 🕒 {created}"):
                st.markdown(f'<div class="result-box">{cont[:350]}…</div>', unsafe_allow_html=True)
                unique_id = _unique_suffix()
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.download_button("📥 نص", cont.encode("utf-8-sig"),
                                       f"اختبار_{sub}.txt", key=f"exam_txt_{eid}_{unique_id}")
                with col2:
                    exam_d = {
                        "school": school_name, "wilaya": wilaya, "grade": gr,
                        "year": school_year, "district": "...", "semester": sem,
                        "subject": sub, "duration": "ساعتان", "content": cont,
                    }
                    epdf = generate_exam_pdf(exam_d)
                    st.download_button("📄 PDF", epdf, f"اختبار_{sub}.pdf",
                                       "application/pdf", key=f"exam_pdf_{eid}_{unique_id}")
                with col3:
                    if _DOCX_AVAILABLE:
                        exam_docx_data = {
                            "school": school_name, "wilaya": wilaya,
                            "grade": gr, "year": school_year,
                            "district": "...", "semester": sem,
                            "subject": sub, "duration": "ساعتان",
                            "content": cont
                        }
                        docx_exam = generate_exam_docx(exam_docx_data)
                        st.download_button("📝 Word", docx_exam, f"اختبار_{sub}.docx",
                                           "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                           key=f"exam_docx_{eid}_{unique_id}")
                    else:
                        st.caption("⚠️ Word غير متاح")
    with arch_tabs[3]:
        corrs = db_exec("SELECT * FROM corrections ORDER BY created_at DESC", fetch=True) or []
        if not corrs:
            st.info("لا توجد تصحيحات.")
        else:
            df_c = pd.DataFrame(corrs,
                                columns=["id", "الاسم", "المادة", "العلامة",
                                         "من", "الملاحظات", "التاريخ"])
            st.dataframe(df_c[["الاسم", "المادة", "العلامة", "من", "التاريخ"]],
                         use_container_width=True)

# ========== TAB 9 — Statistics ==========
with tab_stats:
    total_ex, plans_cnt, exams_cnt, corr_cnt = get_stats()
    st.markdown("### 📉 إحصائيات الاستخدام")
    s1, s2, s3, s4 = st.columns(4)
    for col, val, lbl, clr in [
        (s1, total_ex, "التمارين المولّدة", "#667eea"),
        (s2, plans_cnt, "المذكرات المعدّة", "#764ba2"),
        (s3, exams_cnt, "الاختبارات المولّدة", "#10b981"),
        (s4, corr_cnt, "الأوراق المصحّحة", "#f59e0b"),
    ]:
        with col:
            st.markdown(
                f'<div class="stat-card"><h2 style="color:{clr}">{val}</h2>'
                f'<p>{lbl}</p></div>', unsafe_allow_html=True)
    exercises_all = db_exec("SELECT * FROM exercises ORDER BY created_at DESC", fetch=True) or []
    if exercises_all:
        df_ex = pd.DataFrame(exercises_all,
            columns=["id", "level", "grade", "branch", "subject",
                     "lesson", "ex_type", "difficulty", "content", "created_at"])
        ch1, ch2 = st.columns(2)
        with ch1:
            sc = df_ex["subject"].value_counts().reset_index()
            sc.columns = ["المادة", "العدد"]
            fig_s = px.bar(sc, x="المادة", y="العدد",
                           title="التمارين حسب المادة",
                           template="plotly_dark",
                           color_discrete_sequence=["#667eea"])
            fig_s.update_layout(paper_bgcolor="rgba(0,0,0,0)",
                                 plot_bgcolor="rgba(0,0,0,0)")
            st.plotly_chart(fig_s, use_container_width=True)
        with ch2:
            dc = df_ex["difficulty"].value_counts().reset_index()
            dc.columns = ["الصعوبة", "العدد"]
            fig_d = px.pie(dc, values="العدد", names="الصعوبة",
                           title="توزيع مستويات الصعوبة",
                           template="plotly_dark",
                           color_discrete_sequence=px.colors.qualitative.Pastel)
            fig_d.update_layout(paper_bgcolor="rgba(0,0,0,0)")
            st.plotly_chart(fig_d, use_container_width=True)
    st.markdown("---")
    st.markdown("### ☁️ حالة الربط")
    c1, c2 = st.columns(2)
    with c1:
        if GROQ_API_KEY:
            st.markdown('<div class="success-box">✅ Groq: متصل</div>', unsafe_allow_html=True)
        else:
            st.markdown('<div class="error-box">❌ Groq: غير متصل</div>', unsafe_allow_html=True)
    with c2:
        arcee_connected = test_arcee_connection()
        if arcee_connected:
            st.markdown('<div class="success-box">✅ Arcee: متصل</div>', unsafe_allow_html=True)
        else:
            st.markdown('<div class="error-box">❌ Arcee: غير متصل</div>', unsafe_allow_html=True)

# ========== FOOTER ==========
st.markdown(
    f"""
<div class="donia-ip-footer">
  <div style="margin-bottom:.5rem;font-size:1rem">
    {COPYRIGHT_FOOTER_AR}
  </div>
  <div class="donia-footer-social">
    <a href="{SOCIAL_URL_WHATSAPP}" target="_blank" rel="noopener noreferrer">
      📱 واتساب
    </a>
    <a href="{SOCIAL_URL_FACEBOOK}" target="_blank" rel="noopener noreferrer">
      📖 فيسبوك
    </a>
    <a href="{SOCIAL_URL_TELEGRAM}" target="_blank" rel="noopener noreferrer">
      ✈️ تيليغرام
    </a>
    <a href="{SOCIAL_URL_LINKEDIN}" target="_blank" rel="noopener noreferrer">
      💼 لينكدإن
    </a>
  </div>
  <div style="margin-top:.4rem;font-size:.78rem;color:#888">
    DONIA LABS TECH — منصة المعلم الجزائري الذكي | v5.0 (Dual‑Intelligence + Template Learning + RAG)
  </div>
</div>
""",
    unsafe_allow_html=True,
)
